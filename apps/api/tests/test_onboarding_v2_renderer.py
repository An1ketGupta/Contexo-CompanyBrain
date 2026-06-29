"""Unit tests for the Onboarding v2 PDF renderer — strict Jinja behavior.

The Gotenberg DOCX→PDF conversion isn't tested here (it requires the
sidecar). These tests cover the docxtpl + Jinja2 StrictUndefined path,
which is where production bugs (typos in customer templates) hide.
"""
from __future__ import annotations

import io

import pytest


def _build_simple_docx(body_paragraphs: list[str]) -> bytes:
    """Build a tiny .docx with the given paragraphs as Jinja placeholders.
    docxtpl renders these the same way as a Word-authored document — the
    XML structure is what matters, not how it was authored."""
    from docx import Document

    doc = Document()
    for p in body_paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _extract_text(docx_bytes: bytes) -> str:
    """Open a rendered DOCX (a ZIP) and return concatenated body text.
    Used because docx files are zipped XML — the literal text isn't
    visible in the raw bytes."""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


@pytest.mark.asyncio
async def test_fill_docx_strict_raises_on_unknown_var() -> None:
    """The whole point of strict mode: an unknown {{ var }} must raise
    TemplateVariableError, not silently leave the placeholder."""
    from app.services.pdf import TemplateVariableError, fill_docx_template

    docx = _build_simple_docx(
        ["Hello {{ candidate_name }} — you owe us {{ secret_variable }}."]
    )
    with pytest.raises(TemplateVariableError) as exc_info:
        await fill_docx_template(
            template_bytes=docx,
            context={"candidate_name": "Asha"},
            strict=True,
            template_kind="loi",
        )
    assert exc_info.value.variable_name == "secret_variable"
    assert "secret_variable" in str(exc_info.value)
    assert exc_info.value.template_kind == "loi"


@pytest.mark.asyncio
async def test_fill_docx_strict_passes_when_all_vars_present() -> None:
    from app.services.pdf import fill_docx_template

    docx = _build_simple_docx(
        ["Hello {{ candidate_name }}, welcome to {{ company_name }}."]
    )
    out = await fill_docx_template(
        template_bytes=docx,
        context={"candidate_name": "Asha", "company_name": "Acme"},
        strict=True,
    )
    text = _extract_text(out)
    assert "Asha" in text
    assert "Acme" in text
    assert "{{ candidate_name }}" not in text


@pytest.mark.asyncio
async def test_fill_docx_non_strict_leaves_unknown_placeholders() -> None:
    """Sanity check: non-strict mode is the old behaviour. We don't use
    non-strict in production but it's an escape hatch for diagnostics."""
    from app.services.pdf import fill_docx_template

    docx = _build_simple_docx(["Hello {{ candidate_name }} + {{ missing }}"])
    out = await fill_docx_template(
        template_bytes=docx,
        context={"candidate_name": "Asha"},
        strict=False,
    )
    # Non-strict drops missing — Jinja2 default Undefined renders as ''.
    assert "Asha" in _extract_text(out)


