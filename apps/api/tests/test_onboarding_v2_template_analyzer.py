"""Unit tests for the AI-assisted template analyzer.

The LLM call is mocked — these tests only exercise the deterministic pieces:
text extraction, Jinja detection, applying mappings to a DOCX, and the
phantom-blank guard inside `propose_mappings`.
"""
from __future__ import annotations

import io
from typing import Any

import pytest


def _build_docx(paragraphs: list[str]) -> bytes:
    """Build a small DOCX containing the given paragraphs."""
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _extract_paragraph_text(docx_bytes: bytes) -> list[str]:
    from docx import Document

    return [p.text for p in Document(io.BytesIO(docx_bytes)).paragraphs]


def test_has_jinja_placeholders_detects_braces() -> None:
    from app.services.agents.onboarding_v2.template_analyzer import (
        has_jinja_placeholders,
    )

    with_braces = _build_docx(["Hello {{ candidate_name }}, welcome."])
    without_braces = _build_docx(["Hello _________, welcome."])

    assert has_jinja_placeholders(with_braces) is True
    assert has_jinja_placeholders(without_braces) is False


def test_has_jinja_placeholders_tolerates_filters() -> None:
    """Jinja allows filters like {{ name | upper }} — must still detect them."""
    from app.services.agents.onboarding_v2.template_analyzer import (
        has_jinja_placeholders,
    )

    docx = _build_docx(["Dear {{ candidate_name | upper }},"])
    assert has_jinja_placeholders(docx) is True


def test_extract_text_returns_paragraphs_in_order() -> None:
    from app.services.agents.onboarding_v2.template_analyzer import extract_text

    docx = _build_docx(
        ["First paragraph.", "Second paragraph.", "Third paragraph."]
    )
    text = extract_text(docx)
    assert "First paragraph." in text
    assert "Second paragraph." in text
    assert text.index("First paragraph.") < text.index("Second paragraph.")
    assert text.index("Second paragraph.") < text.index("Third paragraph.")


def test_apply_mappings_replaces_blanks_with_jinja() -> None:
    from app.services.agents.onboarding_v2.template_analyzer import (
        ProposedMapping,
        apply_mappings,
    )

    docx = _build_docx(
        [
            "This Letter of Intent is issued to _____________ for the role of [ROLE].",
            "Start date: ___________.",
        ]
    )
    mappings = [
        ProposedMapping(
            blank_text="_____________",
            variable="candidate_name",
            context_before="issued to",
            context_after="for the",
            confidence="high",
        ),
        ProposedMapping(
            blank_text="[ROLE]",
            variable="role_title",
            context_before="role of",
            context_after=".",
            confidence="high",
        ),
        ProposedMapping(
            blank_text="___________",
            variable="start_date",
            context_before="date:",
            context_after=".",
            confidence="high",
        ),
    ]
    out = apply_mappings(docx_bytes=docx, mappings=mappings)
    paragraphs = _extract_paragraph_text(out)
    assert "{{ candidate_name }}" in paragraphs[0]
    assert "{{ role_title }}" in paragraphs[0]
    assert "{{ start_date }}" in paragraphs[1]
    # No leftover blanks of the patterns we replaced.
    assert "_____________" not in paragraphs[0]
    assert "[ROLE]" not in paragraphs[0]


def test_apply_mappings_raises_when_nothing_matches() -> None:
    """If every proposed blank is a phantom (doesn't appear in the doc), we
    must raise rather than silently no-op."""
    from app.services.agents.onboarding_v2.template_analyzer import (
        ProposedMapping,
        TemplateAnalyzerError,
        apply_mappings,
    )

    docx = _build_docx(["Hello world."])
    mappings = [
        ProposedMapping(
            blank_text="<<MISSING>>",
            variable="candidate_name",
            context_before="",
            context_after="",
            confidence="low",
        )
    ]
    with pytest.raises(TemplateAnalyzerError) as exc_info:
        apply_mappings(docx_bytes=docx, mappings=mappings)
    assert "couldn't apply" in str(exc_info.value).lower()


def test_apply_mappings_handles_multiple_occurrences_of_same_blank() -> None:
    """Same `_________` pattern can appear in multiple places. We replace
    ALL occurrences with the same variable. HR is expected to deduplicate
    if they want different variables in each spot."""
    from app.services.agents.onboarding_v2.template_analyzer import (
        ProposedMapping,
        apply_mappings,
    )

    docx = _build_docx(["First: _________. Second: _________."])
    out = apply_mappings(
        docx_bytes=docx,
        mappings=[
            ProposedMapping(
                blank_text="_________",
                variable="candidate_name",
                context_before="",
                context_after="",
                confidence="high",
            )
        ],
    )
    paragraphs = _extract_paragraph_text(out)
    assert paragraphs[0].count("{{ candidate_name }}") == 2


def test_apply_mappings_preserves_unrelated_text() -> None:
    from app.services.agents.onboarding_v2.template_analyzer import (
        ProposedMapping,
        apply_mappings,
    )

    docx = _build_docx(
        [
            "Confidentiality clause: The Employee shall not disclose any "
            "proprietary information. Candidate: _______.",
        ]
    )
    out = apply_mappings(
        docx_bytes=docx,
        mappings=[
            ProposedMapping(
                blank_text="_______",
                variable="candidate_name",
                context_before="Candidate:",
                context_after=".",
                confidence="high",
            )
        ],
    )
    paragraphs = _extract_paragraph_text(out)
    # Legal text untouched.
    assert "Confidentiality clause" in paragraphs[0]
    assert "proprietary information" in paragraphs[0]
    assert "{{ candidate_name }}" in paragraphs[0]


@pytest.mark.asyncio
async def test_propose_mappings_skips_unknown_variables(monkeypatch: pytest.MonkeyPatch) -> None:
    """The LLM might propose a variable name that isn't in our vocabulary.
    The analyzer must drop those mappings rather than passing them through."""
    from app.services.agents.onboarding_v2 import template_analyzer
    from app.services.llm.types import LLMResponse

    class _FakeClient:
        async def complete(self, **_kwargs: Any) -> LLMResponse:
            payload = (
                '{"mappings": ['
                '{"blank_text": "_____", "variable": "candidate_name", '
                '"context_before": "", "context_after": "", "confidence": "high"},'
                '{"blank_text": "_____", "variable": "made_up_var", '
                '"context_before": "", "context_after": "", "confidence": "low"}'
                "]}"
            )
            return LLMResponse(text=payload)

    monkeypatch.setattr(template_analyzer, "get_llm_client", lambda: _FakeClient())

    result = await template_analyzer.propose_mappings(
        docx_text="Hello _____ welcome.",
        template_kind="loi",
    )
    assert len(result) == 1
    assert result[0].variable == "candidate_name"


@pytest.mark.asyncio
async def test_propose_mappings_skips_phantom_blanks(monkeypatch: pytest.MonkeyPatch) -> None:
    """If the LLM cites a blank that doesn't appear verbatim in the document,
    we drop it — applying would either no-op or corrupt the doc."""
    from app.services.agents.onboarding_v2 import template_analyzer
    from app.services.llm.types import LLMResponse

    class _FakeClient:
        async def complete(self, **_kwargs: Any) -> LLMResponse:
            payload = (
                '{"mappings": ['
                '{"blank_text": "_____", "variable": "candidate_name", '
                '"context_before": "", "context_after": "", "confidence": "high"},'
                '{"blank_text": "<<NOT IN DOC>>", "variable": "role_title", '
                '"context_before": "", "context_after": "", "confidence": "high"}'
                "]}"
            )
            return LLMResponse(text=payload)

    monkeypatch.setattr(template_analyzer, "get_llm_client", lambda: _FakeClient())

    result = await template_analyzer.propose_mappings(
        docx_text="Hello _____ welcome.",
        template_kind="loi",
    )
    assert len(result) == 1
    assert result[0].blank_text == "_____"


@pytest.mark.asyncio
async def test_propose_mappings_tolerates_fenced_json(monkeypatch: pytest.MonkeyPatch) -> None:
    """LLMs occasionally wrap JSON in ```json fences even when instructed not to.
    The analyzer must strip them and still parse."""
    from app.services.agents.onboarding_v2 import template_analyzer
    from app.services.llm.types import LLMResponse

    class _FakeClient:
        async def complete(self, **_kwargs: Any) -> LLMResponse:
            payload = (
                "```json\n"
                '{"mappings": ['
                '{"blank_text": "_____", "variable": "candidate_name", '
                '"context_before": "", "context_after": "", "confidence": "high"}'
                "]}\n"
                "```"
            )
            return LLMResponse(text=payload)

    monkeypatch.setattr(template_analyzer, "get_llm_client", lambda: _FakeClient())

    result = await template_analyzer.propose_mappings(
        docx_text="Hello _____ welcome.",
        template_kind="loi",
    )
    assert len(result) == 1


@pytest.mark.asyncio
async def test_propose_mappings_skips_jinja_regions(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Mixed template safety: if the LLM proposes a mapping whose `blank_text`
    is itself a `{{ var }}` Jinja region, we drop it. Re-templating an
    already-templated spot would corrupt the doc."""
    from app.services.agents.onboarding_v2 import template_analyzer
    from app.services.llm.types import LLMResponse

    class _FakeClient:
        async def complete(self, **_kwargs: Any) -> LLMResponse:
            payload = (
                '{"mappings": ['
                '{"blank_text": "_____", "variable": "candidate_name", '
                '"context_before": "", "context_after": "", "confidence": "high"},'
                # LLM accidentally proposed an already-templated region.
                '{"blank_text": "{{ start_date }}", "variable": "start_date", '
                '"context_before": "", "context_after": "", "confidence": "high"}'
                "]}"
            )
            return LLMResponse(text=payload)

    monkeypatch.setattr(template_analyzer, "get_llm_client", lambda: _FakeClient())

    result = await template_analyzer.propose_mappings(
        docx_text="Hello _____ — start on {{ start_date }}.",
        template_kind="loi",
    )
    assert len(result) == 1
    assert result[0].variable == "candidate_name"


def test_apply_mappings_on_mixed_template_preserves_existing_placeholders() -> None:
    """A template with both existing `{{ var }}` placeholders and plain blanks
    should keep the existing placeholders intact while filling in the blanks."""
    from app.services.agents.onboarding_v2.template_analyzer import (
        ProposedMapping,
        apply_mappings,
    )

    docx = _build_docx(
        [
            "Dear {{ candidate_name }},",
            "We are pleased to offer you the role of _________.",
            "Your CTC is {{ ctc }}.",
            "Start date: ___________.",
        ]
    )
    mappings = [
        ProposedMapping(
            blank_text="_________",
            variable="role_title",
            context_before="role of",
            context_after=".",
            confidence="high",
        ),
        ProposedMapping(
            blank_text="___________",
            variable="start_date",
            context_before="date:",
            context_after=".",
            confidence="high",
        ),
    ]
    out = apply_mappings(docx_bytes=docx, mappings=mappings)
    paragraphs = _extract_paragraph_text(out)
    # Existing placeholders untouched.
    assert "{{ candidate_name }}" in paragraphs[0]
    assert "{{ ctc }}" in paragraphs[2]
    # New placeholders inserted at the blanks.
    assert "{{ role_title }}" in paragraphs[1]
    assert "{{ start_date }}" in paragraphs[3]
    # No leftover plain blanks of the patterns we replaced.
    assert "_________" not in paragraphs[1]
    assert "___________" not in paragraphs[3]


@pytest.mark.asyncio
async def test_propose_mappings_raises_on_unparseable_output(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from app.services.agents.onboarding_v2 import template_analyzer
    from app.services.llm.types import LLMResponse

    class _FakeClient:
        async def complete(self, **_kwargs: Any) -> LLMResponse:
            return LLMResponse(text="this is not JSON at all")

    monkeypatch.setattr(template_analyzer, "get_llm_client", lambda: _FakeClient())

    with pytest.raises(template_analyzer.TemplateAnalyzerError):
        await template_analyzer.propose_mappings(
            docx_text="Hello _____ welcome.",
            template_kind="loi",
        )
