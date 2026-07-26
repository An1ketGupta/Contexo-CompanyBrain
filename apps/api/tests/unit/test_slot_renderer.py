"""Filling a DOCX by position, with no templating engine involved.

What matters here is that the renderer either produces a correct document or
refuses to write at all. A partially-filled offer letter is worse than a failed
run, so both guards (unknown variable, drifted anchor) must fire before any
bytes change.
"""
from __future__ import annotations

import io

import pytest
from docx import Document

from app.services.agents.onboarding_v2.blank_detector import find_all_candidates
from app.services.agents.onboarding_v2.docx_positions import (
    canonical_paragraphs,
    paragraph_hash,
    paragraph_run_text,
)
from app.services.agents.onboarding_v2.template_vars import SIGNATURE_BLOCK_MARKERS
from app.services.pdf.renderer import TemplateVariableError
from app.services.pdf.slot_renderer import SlotDriftError, fill_docx_slots

CONTEXT = {
    "candidate_name": "Aniket Gupta",
    "ctc": "INR 2,000,000.00",
    "start_date": "2026-08-01",
    "role_title": "FullStack Developer",
    **SIGNATURE_BLOCK_MARKERS,
}


def _docx(paragraphs: list[str] | None = None, table: list[list[str]] | None = None) -> bytes:
    doc = Document()
    for p in paragraphs or []:
        doc.add_paragraph(p)
    if table:
        t = doc.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for c, cell in enumerate(row):
                t.cell(r, c).text = cell
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _slots(raw: bytes, variables: dict[int, str]) -> list[dict]:
    """Detect, then attach variables — mirrors what `build_slot_rows` persists."""
    doc = Document(io.BytesIO(raw))
    canonical = canonical_paragraphs(doc)
    out = []
    for c in find_all_candidates(raw):
        variable = variables.get(c.candidate_id)
        if variable is None:
            continue
        para, _kind = canonical[c.paragraph_index]
        out.append(
            {
                "id": f"slot-{c.candidate_id}",
                "paragraph_index": c.paragraph_index,
                "paragraph_hash": paragraph_hash(paragraph_run_text(para)),
                "start_offset": c.start_offset,
                "end_offset": c.end_offset,
                "action": c.action,
                "variable": variable,
                "status": "confirmed",
            }
        )
    return out


def _text(raw: bytes) -> list[str]:
    """Body text by canonical index. Note this keeps blank paragraphs, because
    the indices ARE the addressing scheme — a test that filtered them would stop
    matching the offsets the renderer uses."""
    doc = Document(io.BytesIO(raw))
    return [paragraph_run_text(p) for p, _k in canonical_paragraphs(doc)]


def _nonempty(raw: bytes) -> list[str]:
    """Just the visible lines. Every section contributes an empty header and
    footer paragraph to the canonical enumeration, so a whole-document
    comparison has to ignore blanks."""
    return [t for t in _text(raw) if t.strip()]


# ── The three fill actions ─────────────────────────────────────────────────


def test_replace_span_writes_the_value_over_the_blank() -> None:
    raw = _docx(["Dear [CANDIDATE NAME],"])
    filled = fill_docx_slots(
        docx_bytes=raw, slots=_slots(raw, {0: "candidate_name"}), context=CONTEXT
    )
    assert _text(filled)[0] == "Dear Aniket Gupta,"


def test_insert_after_label_appends_to_an_empty_label_line() -> None:
    """No span to replace — the old pipeline couldn't fill this at all."""
    raw = _docx(["Employee Signature:"])
    filled = fill_docx_slots(
        docx_bytes=raw, slots=_slots(raw, {0: "candidate_name"}), context=CONTEXT
    )
    assert _text(filled)[0] == "Employee Signature: Aniket Gupta"


def test_insert_empty_cell_fills_the_cell_beside_the_label() -> None:
    raw = _docx(table=[["Name:", ""]])
    filled = fill_docx_slots(
        docx_bytes=raw, slots=_slots(raw, {0: "candidate_name"}), context=CONTEXT
    )
    assert _nonempty(filled) == ["Name:", "Aniket Gupta"]


def test_two_blanks_in_one_paragraph_both_land_correctly() -> None:
    """Right-to-left application: filling the first span must not shift the
    second span's recorded offsets."""
    raw = _docx(["Role: [ROLE] starting [START]."])
    filled = fill_docx_slots(
        docx_bytes=raw,
        slots=_slots(raw, {0: "role_title", 1: "start_date"}),
        context=CONTEXT,
    )
    assert _text(filled)[0] == "Role: FullStack Developer starting 2026-08-01."


def test_signature_sentinel_reaches_the_document_unchanged() -> None:
    """apps/esign scans the rendered PDF for this exact string to place the
    signature field — if the renderer mangles it, signing silently falls back
    to a default position."""
    raw = _docx(["Employee Signature:"])
    filled = fill_docx_slots(
        docx_bytes=raw, slots=_slots(raw, {0: "hr_signature_block"}), context=CONTEXT
    )
    assert SIGNATURE_BLOCK_MARKERS["hr_signature_block"] in _text(filled)[0]


def test_no_slots_returns_the_document_untouched() -> None:
    raw = _docx(["Nothing to fill here."])
    assert fill_docx_slots(docx_bytes=raw, slots=[], context=CONTEXT) == raw


# ── Guards: refuse rather than write something wrong ───────────────────────


def test_unknown_variable_raises_before_writing_anything() -> None:
    raw = _docx(["Dear [CANDIDATE NAME], your CTC is ______."])
    slots = _slots(raw, {0: "candidate_name", 1: "ctc"})
    with pytest.raises(TemplateVariableError) as exc:
        fill_docx_slots(
            docx_bytes=raw, slots=slots, context={"candidate_name": "Aniket Gupta"}
        )
    assert exc.value.variable_name == "ctc"


def test_confirmed_slot_with_no_variable_is_rejected() -> None:
    """The DB forbids this too — a confirmed field with nothing to fill it would
    render as an empty string in a contract."""
    raw = _docx(["Dear [CANDIDATE NAME],"])
    slots = _slots(raw, {0: "candidate_name"})
    slots[0]["variable"] = None
    with pytest.raises(Exception, match="no variable assigned"):
        fill_docx_slots(docx_bytes=raw, slots=slots, context=CONTEXT)


def test_editing_a_mapped_paragraph_raises_drift() -> None:
    raw = _docx(["Dear [CANDIDATE NAME],"])
    slots = _slots(raw, {0: "candidate_name"})

    doc = Document(io.BytesIO(raw))
    canonical_paragraphs(doc)[0][0].runs[0].text = "Dear [CANDIDATE NAME], (revised)"
    buf = io.BytesIO()
    doc.save(buf)

    with pytest.raises(SlotDriftError) as exc:
        fill_docx_slots(docx_bytes=buf.getvalue(), slots=slots, context=CONTEXT)
    assert exc.value.paragraph_indexes == [0]


def test_editing_an_unrelated_paragraph_does_not_invalidate_other_fields() -> None:
    """The reason drift is checked per-paragraph and not per-document: HR
    rewording one clause of a six-page letter must not force them to re-confirm
    every field in it."""
    raw = _docx(
        [
            "Dear [CANDIDATE NAME],",
            "Some boilerplate clause.",
            "Your CTC will be ______ per annum.",
        ]
    )
    slots = _slots(raw, {0: "candidate_name", 1: "ctc"})

    doc = Document(io.BytesIO(raw))
    canonical_paragraphs(doc)[1][0].runs[0].text = "Some boilerplate clause, reworded."
    buf = io.BytesIO()
    doc.save(buf)

    filled = fill_docx_slots(docx_bytes=buf.getvalue(), slots=slots, context=CONTEXT)
    text = _text(filled)
    assert text[0] == "Dear Aniket Gupta,"
    assert text[2] == "Your CTC will be INR 2,000,000.00 per annum."


def test_deleted_paragraph_counts_as_drift_not_a_crash() -> None:
    raw = _docx(["Dear [CANDIDATE NAME],", "Second paragraph."])
    slots = _slots(raw, {0: "candidate_name"})
    slots[0]["paragraph_index"] = 99  # anchor that no longer resolves

    with pytest.raises(SlotDriftError):
        fill_docx_slots(docx_bytes=raw, slots=slots, context=CONTEXT)


# ── Render context consolidation ───────────────────────────────────────────


def test_sample_context_supplies_every_vocabulary_variable() -> None:
    """The classifier may map a blank to any vocabulary variable, so every one
    of them must exist in the render context — otherwise a valid mapping blows
    up as TemplateVariableError during generation."""
    from app.services.agents.onboarding_v2.render_context import missing_vocabulary_keys

    assert missing_vocabulary_keys() == []
