"""Detection heuristics for HR template fill-points.

These tests encode the two failures that motivated replacing the old pipeline:

  * a label with nothing after it, and an empty table cell beside a labelled
    one, were invisible to a marker-span regex — no candidate, no error, and a
    contract that shipped with the field blank;
  * and the guard against over-correcting: a prose clause ending in a colon
    ("The terms are as follows:") must never be treated as a field, or we'd
    splice a candidate's salary into the middle of a sentence.
"""
from __future__ import annotations

import io

import pytest
from docx import Document

from app.services.agents.onboarding_v2.blank_detector import (
    ACTION_INSERT_AFTER_LABEL,
    ACTION_INSERT_EMPTY_CELL,
    ACTION_REPLACE_SPAN,
    find_all_candidates,
    find_blank_candidates,
    find_empty_cell_candidates,
    find_label_colon_candidates,
    scan_for_unfilled_signals,
)


def _docx(paragraphs: list[str] | None = None, table: list[list[str]] | None = None) -> bytes:
    doc = Document()
    for p in paragraphs or []:
        doc.add_paragraph(p)
    if table:
        t = doc.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for c, cell in enumerate(row):
                t.cell(r, c).text = cell
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Heuristic 1: marker spans (the original detector) ──────────────────────


@pytest.mark.parametrize(
    "text,expected",
    [
        ("Your CTC will be ____________ per annum.", "____________"),
        ("Dear [CANDIDATE NAME],", "[CANDIDATE NAME]"),
        ("Dear <NAME>,", "<NAME>"),
        ("Signed on ......... at", "........."),
        ("Hello {name} there", "{name}"),
    ],
)
def test_marker_spans_detected(text: str, expected: str) -> None:
    found = find_blank_candidates(_docx([text]))
    assert [c.matched_text for c in found] == [expected]
    assert found[0].action == ACTION_REPLACE_SPAN


def test_existing_jinja_placeholder_is_not_a_blank() -> None:
    """A real `{{ var }}` is already mapped — re-templating it would nest tags."""
    assert find_blank_candidates(_docx(["Dear {{ candidate_name }},"])) == []


def test_offsets_slice_the_exact_blank() -> None:
    """Offsets are the contract between detection and rendering: everything
    downstream substitutes by position, never by re-finding the text."""
    text = "Your CTC will be ____________ per annum."
    c = find_blank_candidates(_docx([text]))[0]
    assert text[c.start_offset:c.end_offset] == "____________"


def test_two_blanks_in_one_paragraph_are_distinguished_by_position() -> None:
    found = find_blank_candidates(_docx(["Role: [ROLE] starting [START]."]))
    assert len(found) == 2
    assert found[0].start_offset < found[1].start_offset


# ── Heuristic 2: a label with nothing after it ─────────────────────────────


def test_label_with_nothing_after_is_detected() -> None:
    """The first silent-failure case: no marker characters at all, so the old
    regex-only detector produced no candidate and the field shipped blank."""
    found = find_label_colon_candidates(_docx(["Employee Signature:"]))
    assert len(found) == 1
    assert found[0].action == ACTION_INSERT_AFTER_LABEL
    # An insertion point, not a span to replace.
    assert found[0].start_offset == found[0].end_offset == len("Employee Signature:")
    assert found[0].matched_text == ""


@pytest.mark.parametrize(
    "text",
    [
        "The terms are as follows:",
        "Your compensation will be as under:",
        "We hereby confirm the following:",
        "This agreement includes:",
        "NOW THEREFORE the parties agree as follows:",
    ],
)
def test_prose_lead_ins_are_not_labels(text: str) -> None:
    """False positives here are worse than misses — a value injected after a
    clause corrupts the sentence in a legally-binding document."""
    assert find_label_colon_candidates(_docx([text])) == []


@pytest.mark.parametrize(
    "text",
    ["Name:", "Date of Joining:", "Name of the Employee:", "CTC:"],
)
def test_real_field_labels_are_detected(text: str) -> None:
    assert len(find_label_colon_candidates(_docx([text]))) == 1


def test_label_that_already_has_a_blank_is_not_double_counted() -> None:
    """`Name: ______` is one field, found by the span heuristic — it must not
    also produce an insert-after-label candidate at the end of the line."""
    assert find_label_colon_candidates(_docx(["Name: ______"])) == []


# ── Heuristic 3: empty table cell beside a labelled one ────────────────────


def test_empty_cell_beside_label_is_detected() -> None:
    """The second silent-failure case: the fill-point is an absence of text."""
    found = find_empty_cell_candidates(_docx(table=[["Name:", ""]]))
    assert len(found) == 1
    assert found[0].action == ACTION_INSERT_EMPTY_CELL
    assert "Name:" in found[0].context_before


def test_filled_cell_is_not_a_fill_point() -> None:
    assert find_empty_cell_candidates(_docx(table=[["Designation", "Engineer"]])) == []


def test_empty_cell_without_a_label_is_ignored() -> None:
    """No adjacent label means nothing to infer a variable from — flagging it
    would just produce a field HR has to dismiss."""
    assert find_empty_cell_candidates(_docx(table=[["", ""]])) == []


def test_label_word_without_colon_still_counts() -> None:
    found = find_empty_cell_candidates(_docx(table=[["Designation", ""]]))
    assert len(found) == 1


# ── Combined ───────────────────────────────────────────────────────────────


def test_find_all_candidates_covers_every_shape_once() -> None:
    raw = _docx(
        paragraphs=[
            "Dear [CANDIDATE NAME],",
            "Your CTC will be ____________ per annum.",
            "The terms are as follows:",
            "Employee Signature:",
        ],
        table=[["Name:", ""]],
    )
    found = find_all_candidates(raw)
    actions = sorted(c.action for c in found)
    assert actions == sorted(
        [
            ACTION_REPLACE_SPAN,
            ACTION_REPLACE_SPAN,
            ACTION_INSERT_AFTER_LABEL,
            ACTION_INSERT_EMPTY_CELL,
        ]
    )
    # candidate_ids must be dense and unique — the classifier keys on them.
    assert sorted(c.candidate_id for c in found) == list(range(len(found)))


def test_labelled_cell_does_not_compete_with_its_empty_neighbour() -> None:
    """`Name: | <empty>` is ONE field. Without suppression the label cell would
    also match the label-colon heuristic and produce a second, conflicting
    fill-point writing to a different spot."""
    found = find_all_candidates(_docx(table=[["Name:", ""]]))
    assert len(found) == 1
    assert found[0].action == ACTION_INSERT_EMPTY_CELL


# ── Post-render safety net ─────────────────────────────────────────────────


def test_scan_flags_output_that_still_looks_unfilled() -> None:
    warnings = scan_for_unfilled_signals(
        "Your CTC will be ________ per annum.\nEmployee Signature:\n"
    )
    assert len(warnings) == 2


def test_scan_is_quiet_on_a_fully_filled_document() -> None:
    assert scan_for_unfilled_signals(
        "Your CTC will be INR 2,000,000.00 per annum.\n"
        "Employee Signature: Aniket Gupta\n"
        "The terms are as follows:\n"
    ) == []
