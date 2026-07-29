"""Run-level splicing tests.

`test_bold_label_does_not_bleed_into_the_value` is the regression that motivated
this module: the previous renderer collapsed every run of an edited paragraph
into `runs[0]`, so a bold label smeared its formatting across the value beside
it. These tests assert the runs directly rather than the visible text, because
the text was never the thing that was broken.
"""
from __future__ import annotations

import io

import pytest
from docx import Document
from docx.shared import Pt

from app.services.documents.docx_positions import paragraph_run_text
from app.services.documents.generation.docx_splice import SpliceEdit, splice_paragraph


def _para(runs: list[tuple[str, dict]]):
    """Build a one-paragraph document from (text, formatting) pairs."""
    doc = Document()
    p = doc.add_paragraph()
    for text, fmt in runs:
        r = p.add_run(text)
        for key, value in fmt.items():
            if key == "size":
                r.font.size = Pt(value)
            elif key == "name":
                r.font.name = value
            else:
                setattr(r, key, value)
    return doc, p


def _roundtrip(doc):
    """Save and reopen, so assertions run against real serialized XML."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf)


# ── The regression ────────────────────────────────────────────────────────


def test_bold_label_does_not_bleed_into_the_value():
    doc, p = _para([("Annual Salary: ", {"bold": True}), ("₹12,00,000", {})])
    text = paragraph_run_text(p)
    start = text.index("₹12,00,000")

    splice_paragraph(p, [SpliceEdit(start, start + len("₹12,00,000"), "₹18,50,000")])

    reopened = _roundtrip(doc).paragraphs[0]
    assert paragraph_run_text(reopened) == "Annual Salary: ₹18,50,000"

    label, value = reopened.runs[0], reopened.runs[1]
    assert label.text == "Annual Salary: "
    assert label.bold is True, "the label must stay bold"
    assert value.text == "₹18,50,000"
    assert not value.bold, "the value must NOT inherit the label's bold"


def test_untouched_runs_are_not_reassigned_at_all():
    """A run that overlaps no edit must come out byte-identical, including its
    font, size, and italics."""
    doc, p = _para([
        ("Reporting to ", {}),
        ("John Smith", {"italic": True, "name": "Georgia", "size": 14}),
        (" in ", {}),
        ("Bengaluru", {"bold": True}),
    ])
    text = paragraph_run_text(p)
    start = text.index("Bengaluru")

    splice_paragraph(p, [SpliceEdit(start, start + len("Bengaluru"), "Mumbai")])

    runs = _roundtrip(doc).paragraphs[0].runs
    manager = runs[1]
    assert manager.text == "John Smith"
    assert manager.italic is True
    assert manager.font.name == "Georgia"
    assert manager.font.size == Pt(14)
    assert runs[3].text == "Mumbai"
    assert runs[3].bold is True


def test_many_fields_in_one_paragraph_each_keep_their_own_style():
    doc, p = _para([
        ("Name: ", {"bold": True}),
        ("Rahul", {}),
        (" | Date: ", {"bold": True}),
        ("15 Aug 2026", {"italic": True}),
    ])
    text = paragraph_run_text(p)

    splice_paragraph(p, [
        SpliceEdit(text.index("Rahul"), text.index("Rahul") + 5, "Priya"),
        SpliceEdit(
            text.index("15 Aug 2026"),
            text.index("15 Aug 2026") + len("15 Aug 2026"),
            "01 Sep 2026",
        ),
    ])

    runs = _roundtrip(doc).paragraphs[0].runs
    assert [r.text for r in runs] == ["Name: ", "Priya", " | Date: ", "01 Sep 2026"]
    assert runs[0].bold is True
    assert not runs[1].bold
    assert runs[2].bold is True
    assert runs[3].italic is True


# ── Span mechanics ────────────────────────────────────────────────────────


def test_replacement_inside_a_single_run_keeps_its_neighbours():
    doc, p = _para([("Dear ", {}), ("Rahul", {}), (", welcome.", {})])
    splice_paragraph(p, [SpliceEdit(5, 10, "Priya Sharma")])
    assert paragraph_run_text(p) == "Dear Priya Sharma, welcome."
    assert [r.text for r in p.runs] == ["Dear ", "Priya Sharma", ", welcome."]


def test_span_crossing_several_runs_collapses_only_that_span():
    """The value takes the style of the run where the span starts; the tail of
    the last crossed run survives."""
    doc, p = _para([
        ("Salary is ", {}),
        ("twelve", {"bold": True}),
        (" lakh", {"italic": True}),
        (" per annum.", {}),
    ])
    text = paragraph_run_text(p)
    start = text.index("twelve")
    end = text.index(" per annum.")

    splice_paragraph(p, [SpliceEdit(start, end, "₹12,00,000")])

    assert paragraph_run_text(p) == "Salary is ₹12,00,000 per annum."
    assert p.runs[0].text == "Salary is "
    assert p.runs[1].text == "₹12,00,000"
    assert p.runs[3].text == " per annum."


def test_partial_overlap_keeps_the_surviving_tail():
    doc, p = _para([("ABCDEF", {}), ("GHIJKL", {})])
    splice_paragraph(p, [SpliceEdit(3, 9, "-")])
    assert paragraph_run_text(p) == "ABC-JKL"


# ── Insertion points ──────────────────────────────────────────────────────


def test_insert_at_end_appends_with_the_last_run_style():
    doc, p = _para([("Employee Signature:", {"bold": True})])
    end = len(paragraph_run_text(p))

    splice_paragraph(p, [SpliceEdit(end, end, " Rahul")])

    reopened = _roundtrip(doc).paragraphs[0]
    assert paragraph_run_text(reopened) == "Employee Signature: Rahul"
    assert reopened.runs[0].bold is True


def test_insert_into_an_empty_paragraph_creates_a_run():
    doc = Document()
    p = doc.add_paragraph()
    assert not p.runs

    splice_paragraph(p, [SpliceEdit(0, 0, "Software Engineer")])
    assert paragraph_run_text(p) == "Software Engineer"


def test_insert_into_a_paragraph_whose_only_run_is_empty():
    doc, p = _para([("", {"bold": True})])
    splice_paragraph(p, [SpliceEdit(0, 0, "Bengaluru")])
    assert paragraph_run_text(p) == "Bengaluru"
    assert p.runs[0].bold is True


def test_insert_in_the_middle_splits_correctly():
    doc, p = _para([("Name:  (signed)", {})])
    splice_paragraph(p, [SpliceEdit(6, 6, "Rahul")])
    assert paragraph_run_text(p) == "Name: Rahul (signed)"


# ── Value handling ────────────────────────────────────────────────────────


def test_empty_replacement_deletes_the_span():
    doc, p = _para([("Name: ", {}), ("PLACEHOLDER", {})])
    splice_paragraph(p, [SpliceEdit(6, 17, "")])
    assert paragraph_run_text(p) == "Name: "


def test_longer_and_shorter_values_both_work():
    for value in ("X", "a much longer replacement value indeed"):
        doc, p = _para([("Role: ", {}), ("Engineer", {})])
        splice_paragraph(p, [SpliceEdit(6, 14, value)])
        assert paragraph_run_text(p) == f"Role: {value}"


def test_unicode_and_currency_survive():
    doc, p = _para([("Salary: ", {}), ("____", {})])
    splice_paragraph(p, [SpliceEdit(8, 12, "₹12,00,000 (बारह लाख)")])
    assert paragraph_run_text(p) == "Salary: ₹12,00,000 (बारह लाख)"


# ── Contract ──────────────────────────────────────────────────────────────


def test_no_edits_is_a_no_op():
    doc, p = _para([("Unchanged", {"bold": True})])
    assert splice_paragraph(p, []) is False
    assert paragraph_run_text(p) == "Unchanged"


def test_returns_false_when_nothing_actually_changed():
    doc, p = _para([("Rahul", {})])
    assert splice_paragraph(p, [SpliceEdit(0, 5, "Rahul")]) is False


def test_returns_true_when_something_changed():
    doc, p = _para([("Rahul", {})])
    assert splice_paragraph(p, [SpliceEdit(0, 5, "Priya")]) is True


def test_edits_applied_out_of_order_land_correctly():
    """Callers may hand edits in any order; offsets are all in the ORIGINAL
    coordinate system and must not shift as earlier edits are applied."""
    doc, p = _para([("A__B__C", {})])
    splice_paragraph(p, [
        SpliceEdit(4, 6, "YY"),   # deliberately the later span first
        SpliceEdit(1, 3, "XXXX"),
    ])
    assert paragraph_run_text(p) == "AXXXXBYYC"


def test_headers_and_table_cells_use_the_same_path():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    cell = table.rows[0].cells[1]
    p = cell.paragraphs[0]
    p.add_run("Old Value").bold = True

    splice_paragraph(p, [SpliceEdit(0, 9, "New Value")])

    reopened = _roundtrip(doc).tables[0].rows[0].cells[1].paragraphs[0]
    assert paragraph_run_text(reopened) == "New Value"
    assert reopened.runs[0].bold is True
