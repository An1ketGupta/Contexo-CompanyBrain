"""Generation renderer tests.

Covers the guards that stand between a template and a legally-binding PDF:
drift detection, missing values, and unmapped fields — all of which must fire
BEFORE any bytes are modified, so a failed render can never leave a
half-completed document.

`test_generated_document_preserves_formatting` is the end-to-end version of the
regression in `test_docx_splice.py`: it runs the real fill path over a real
template and asserts the runs survived.
"""
from __future__ import annotations

import io

import pytest
from docx import Document
from docx.shared import Pt

from app.services.documents.constants import (
    ACTION_INSERT_AFTER_LABEL,
    ACTION_INSERT_EMPTY_CELL,
    ACTION_REPLACE_SPAN,
)
from app.services.documents.docx_positions import (
    canonical_paragraphs,
    paragraph_hash,
    paragraph_run_text,
)
from app.services.documents.generation import renderer as renderer_mod
from app.services.documents.generation.renderer import (
    MissingValueError,
    SlotDriftError,
    UnmappedSlotError,
    fill_docx,
    format_value,
    render_to_pdf,
)


def _template(lines: list[tuple[str, dict]]) -> bytes:
    """One paragraph per (text, formatting) pair."""
    doc = Document()
    for text, fmt in lines:
        p = doc.add_paragraph()
        run = p.add_run(text)
        for key, value in fmt.items():
            if key == "size":
                run.font.size = Pt(value)
            else:
                setattr(run, key, value)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _slot(docx_bytes: bytes, para_idx: int, start: int, end: int, variable: str, **kw):
    """Build a slot row with a genuine paragraph_hash for the given template."""
    doc = Document(io.BytesIO(docx_bytes))
    paragraph, _kind = canonical_paragraphs(doc)[para_idx]
    row = {
        "paragraph_index": para_idx,
        "paragraph_hash": paragraph_hash(paragraph_run_text(paragraph)),
        "start_offset": start,
        "end_offset": end,
        "action": ACTION_REPLACE_SPAN,
        "variable": variable,
    }
    row.update(kw)
    return row


def _text_of(docx_bytes: bytes) -> list[str]:
    doc = Document(io.BytesIO(docx_bytes))
    return [paragraph_run_text(p) for p, _k in canonical_paragraphs(doc)]


# ── The end-to-end formatting guarantee ───────────────────────────────────


def test_generated_document_preserves_formatting():
    template = _template([
        ("Offer of Employment", {"bold": True, "size": 16}),
        ("Annual Salary: ", {"bold": True}),
    ])
    # Append an unstyled value run to paragraph 1.
    doc = Document(io.BytesIO(template))
    doc.paragraphs[1].add_run("Rs 12,00,000")
    buf = io.BytesIO()
    doc.save(buf)
    template = buf.getvalue()

    slots = [_slot(template, 1, 15, 27, "ctc")]
    filled = fill_docx(docx_bytes=template, slots=slots, context={"ctc": "Rs 18,50,000"})

    out = Document(io.BytesIO(filled))
    heading, salary = out.paragraphs[0], out.paragraphs[1]

    assert heading.runs[0].bold is True
    assert heading.runs[0].font.size == Pt(16)
    assert salary.runs[0].text == "Annual Salary: "
    assert salary.runs[0].bold is True
    assert salary.runs[1].text == "Rs 18,50,000"
    assert not salary.runs[1].bold


def test_the_original_template_bytes_are_never_mutated():
    template = _template([("Name: Rahul", {})])
    before = bytes(template)
    fill_docx(
        docx_bytes=template,
        slots=[_slot(template, 0, 6, 11, "name")],
        context={"name": "Priya"},
    )
    assert template == before


# ── Basic filling ─────────────────────────────────────────────────────────


def test_replace_span_writes_the_value():
    template = _template([("Name: Rahul", {})])
    filled = fill_docx(
        docx_bytes=template,
        slots=[_slot(template, 0, 6, 11, "name")],
        context={"name": "Priya Sharma"},
    )
    assert _text_of(filled)[0] == "Name: Priya Sharma"


def test_insert_after_label_adds_a_separator():
    template = _template([("Employee Signature:", {})])
    slot = _slot(template, 0, 19, 19, "signature", action=ACTION_INSERT_AFTER_LABEL)
    filled = fill_docx(docx_bytes=template, slots=[slot], context={"signature": "Rahul"})
    assert _text_of(filled)[0] == "Employee Signature: Rahul"


def test_insert_after_label_does_not_double_the_space():
    template = _template([("Employee Signature: ", {})])
    slot = _slot(template, 0, 20, 20, "signature", action=ACTION_INSERT_AFTER_LABEL)
    filled = fill_docx(docx_bytes=template, slots=[slot], context={"signature": "Rahul"})
    assert _text_of(filled)[0] == "Employee Signature: Rahul"


def test_insert_empty_cell_writes_into_the_cell():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Designation"
    buf = io.BytesIO()
    doc.save(buf)
    template = buf.getvalue()

    canonical = canonical_paragraphs(Document(io.BytesIO(template)))
    target = next(
        i for i, (p, kind) in enumerate(canonical)
        if kind == "table" and not paragraph_run_text(p)
    )
    slot = _slot(template, target, 0, 0, "designation", action=ACTION_INSERT_EMPTY_CELL)

    filled = fill_docx(
        docx_bytes=template, slots=[slot], context={"designation": "Software Engineer"}
    )
    assert "Software Engineer" in _text_of(filled)


def test_several_slots_across_several_paragraphs():
    template = _template([
        ("Dear Rahul,", {}),
        ("Joining Date: 15 August 2026", {}),
        ("Reporting to John Smith.", {}),
    ])
    slots = [
        _slot(template, 0, 5, 10, "name"),
        _slot(template, 1, 14, 28, "start_date"),
        _slot(template, 2, 13, 23, "manager"),
    ]
    filled = fill_docx(
        docx_bytes=template,
        slots=slots,
        context={"name": "Priya", "start_date": "01 September 2026", "manager": "Asha Rao"},
    )
    # Body paragraphs only: canonical order appends the (empty) header and
    # footer paragraphs after the body.
    assert _text_of(filled)[:3] == [
        "Dear Priya,",
        "Joining Date: 01 September 2026",
        "Reporting to Asha Rao.",
    ]


def test_no_slots_passes_the_document_through_untouched():
    template = _template([("Nothing to fill.", {})])
    assert fill_docx(docx_bytes=template, slots=[], context={}) == template


# ── Guards ────────────────────────────────────────────────────────────────


def test_edited_template_raises_drift_and_writes_nothing():
    template = _template([("Name: Rahul", {})])
    slots = [_slot(template, 0, 6, 11, "name")]

    # HR rewords the paragraph after confirming the field.
    edited = _template([("Candidate Name: Rahul", {})])

    with pytest.raises(SlotDriftError) as exc:
        fill_docx(docx_bytes=edited, slots=slots, context={"name": "Priya"})
    assert exc.value.paragraph_indexes == [0]
    assert "confirm the fields again" in str(exc.value)


def test_editing_an_unrelated_paragraph_does_not_trigger_drift():
    """The drift guard is per-paragraph on purpose: rewording a clause
    elsewhere in a long template must not invalidate a confirmed field."""
    template = _template([("Name: Rahul", {}), ("Some boilerplate clause.", {})])
    slots = [_slot(template, 0, 6, 11, "name")]

    edited = _template([("Name: Rahul", {}), ("A completely rewritten clause.", {})])
    filled = fill_docx(docx_bytes=edited, slots=slots, context={"name": "Priya"})
    assert _text_of(filled)[0] == "Name: Priya"


def test_out_of_range_paragraph_index_is_drift_not_a_crash():
    template = _template([("Name: Rahul", {})])
    slot = _slot(template, 0, 6, 11, "name")
    slot["paragraph_index"] = 99
    with pytest.raises(SlotDriftError):
        fill_docx(docx_bytes=template, slots=[slot], context={"name": "Priya"})


def test_missing_value_raises_before_writing():
    template = _template([("Name: Rahul", {})])
    with pytest.raises(MissingValueError) as exc:
        fill_docx(docx_bytes=template, slots=[_slot(template, 0, 6, 11, "name")], context={})
    assert exc.value.variable == "name"


def test_unmapped_slot_raises():
    template = _template([("Name: Rahul", {})])
    with pytest.raises(UnmappedSlotError):
        fill_docx(
            docx_bytes=template,
            slots=[_slot(template, 0, 6, 11, None)],
            context={"name": "Priya"},
        )


def test_explicit_none_value_is_allowed_and_renders_blank():
    """None is a supplied value, not a missing one — an optional field that
    resolved to nothing must render as a blank, not as the text 'None'."""
    template = _template([("Name: Rahul", {})])
    filled = fill_docx(
        docx_bytes=template,
        slots=[_slot(template, 0, 6, 11, "name")],
        context={"name": None},
    )
    assert _text_of(filled)[0] == "Name: "


# ── Value formatting ──────────────────────────────────────────────────────


@pytest.mark.parametrize(
    ("value", "expected"),
    [(None, ""), (True, "Yes"), (False, "No"), (12, "12"), ("x", "x"), (1.5, "1.5")],
)
def test_format_value(value, expected):
    assert format_value(value) == expected


# ── PDF conversion ────────────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_pdf_failure_still_returns_the_docx(monkeypatch):
    """The spec is explicit: a failed PDF conversion must not cost HR the
    generated DOCX."""
    async def _boom(_docx):
        raise RuntimeError("gotenberg unreachable")

    monkeypatch.setattr(renderer_mod, "convert_docx_to_pdf", _boom)

    template = _template([("Name: Rahul", {})])
    filled, pdf = await render_to_pdf(
        docx_bytes=template,
        slots=[_slot(template, 0, 6, 11, "name")],
        context={"name": "Priya"},
    )
    assert pdf is None
    assert _text_of(filled)[0] == "Name: Priya"


@pytest.mark.asyncio
async def test_successful_pdf_conversion_returns_both(monkeypatch):
    async def _fake(_docx):
        return b"%PDF-1.7 fake"

    monkeypatch.setattr(renderer_mod, "convert_docx_to_pdf", _fake)

    template = _template([("Name: Rahul", {})])
    filled, pdf = await render_to_pdf(
        docx_bytes=template,
        slots=[_slot(template, 0, 6, 11, "name")],
        context={"name": "Priya"},
    )
    assert pdf == b"%PDF-1.7 fake"
    assert _text_of(filled)[0] == "Name: Priya"


@pytest.mark.asyncio
async def test_drift_propagates_out_of_render_to_pdf(monkeypatch):
    async def _fake(_docx):
        return b"%PDF"

    monkeypatch.setattr(renderer_mod, "convert_docx_to_pdf", _fake)

    template = _template([("Name: Rahul", {})])
    slots = [_slot(template, 0, 6, 11, "name")]
    edited = _template([("Candidate Name: Rahul", {})])

    with pytest.raises(SlotDriftError):
        await render_to_pdf(docx_bytes=edited, slots=slots, context={"name": "Priya"})
