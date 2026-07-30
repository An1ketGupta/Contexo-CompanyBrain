"""In-place text editing of a generated document.

The behaviour these lock down is the difference between this editor and the
flat-text one it replaces: editing line 3 must not touch line 4's runs. The
old editor rewrote every paragraph on every save, so one typo fix in the
salutation flattened the formatting of the whole document.
"""
from __future__ import annotations

import io

import pytest
from docx import Document
from docx.shared import Pt

from app.services.documents.text_edit import (
    TextEditError,
    apply_paragraph_edits,
    draft_fingerprint,
    extract_editable_paragraphs,
)


def _docx(*paragraphs: str) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    return _bytes(doc)


def _bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _texts(docx_bytes: bytes) -> list[str]:
    return [p.text for p in extract_editable_paragraphs(docx_bytes)]


# ── Reading ───────────────────────────────────────────────────────────────


def test_extracts_non_blank_paragraphs_in_order():
    data = _docx("Dear Rahul,", "We are pleased to offer you the role.", "Regards,")
    paragraphs = extract_editable_paragraphs(data)

    assert [p.text for p in paragraphs] == [
        "Dear Rahul,",
        "We are pleased to offer you the role.",
        "Regards,",
    ]
    assert [p.kind for p in paragraphs] == ["body", "body", "body"]


def test_blank_paragraphs_are_dropped_but_indices_do_not_shift():
    """The index is the document's own numbering — an edit posted for index 2
    must land on the third paragraph, not the second non-blank one."""
    data = _docx("First line", "", "   ", "Fourth line")
    paragraphs = extract_editable_paragraphs(data)

    assert [(p.index, p.text) for p in paragraphs] == [
        (0, "First line"),
        (3, "Fourth line"),
    ]

    edited, changed = apply_paragraph_edits(
        docx_bytes=data, edits={3: "Fourth line, corrected"}
    )
    assert changed == 1
    assert _texts(edited) == ["First line", "Fourth line, corrected"]


def test_table_and_header_paragraphs_are_labelled():
    doc = Document()
    doc.add_paragraph("Body line")
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].paragraphs[0].add_run("Cell line")
    doc.sections[0].header.paragraphs[0].add_run("Header line")

    kinds = {p.text: p.kind for p in extract_editable_paragraphs(_bytes(doc))}
    assert kinds == {
        "Body line": "body",
        "Cell line": "table",
        "Header line": "header",
    }


def test_unreadable_file_is_a_clean_error_not_a_crash():
    with pytest.raises(TextEditError) as exc:
        extract_editable_paragraphs(b"this is not a docx")
    assert "corrupted" in str(exc.value)


# ── Writing ───────────────────────────────────────────────────────────────


def test_editing_one_line_leaves_the_others_untouched():
    data = _docx("Dear Rahul,", "Your start date is 1 August 2026.", "Regards,")

    edited, changed = apply_paragraph_edits(
        docx_bytes=data, edits={1: "Your start date is 15 August 2026."}
    )

    assert changed == 1
    assert _texts(edited) == [
        "Dear Rahul,",
        "Your start date is 15 August 2026.",
        "Regards,",
    ]


def test_untouched_paragraphs_keep_their_run_formatting():
    """The regression the old flat-text editor had: saving a fix to one line
    rewrote every paragraph, flattening styled runs everywhere."""
    doc = Document()
    intro = doc.add_paragraph()
    intro.add_run("Dear ")
    intro.add_run("Rahul Sharma").bold = True
    salary = doc.add_paragraph()
    salary.add_run("Annual CTC: ").bold = True
    run = salary.add_run("₹12,00,000")
    run.italic = True
    run.font.size = Pt(13)

    edited, changed = apply_paragraph_edits(
        docx_bytes=_bytes(doc), edits={0: "Dear Rahul Sharma,"}
    )
    assert changed == 1

    reopened = Document(io.BytesIO(edited))
    untouched = reopened.paragraphs[1].runs
    assert [r.text for r in untouched] == ["Annual CTC: ", "₹12,00,000"]
    assert untouched[0].bold is True
    assert untouched[1].italic is True
    assert untouched[1].font.size == Pt(13)
    assert not untouched[1].bold, "the label's bold must not have bled across"


def test_an_edited_paragraph_keeps_its_first_runs_style():
    """Editing a line does flatten it — a flat string has no run boundaries to
    restore. It must at least keep the paragraph's own styling, not the
    document default."""
    doc = Document()
    p = doc.add_paragraph()
    heading = p.add_run("LOI")
    heading.bold = True
    heading.font.size = Pt(16)

    edited, _ = apply_paragraph_edits(
        docx_bytes=_bytes(doc), edits={0: "LOI (REVISED)"}
    )

    run = Document(io.BytesIO(edited)).paragraphs[0].runs[0]
    assert run.text == "LOI (REVISED)"
    assert run.bold is True
    assert run.font.size == Pt(16)


def test_editing_a_table_cell_works_through_the_same_path():
    doc = Document()
    doc.add_paragraph("Compensation")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("Basic")
    table.rows[0].cells[1].paragraphs[0].add_run("₹6,00,000")

    edited, changed = apply_paragraph_edits(
        docx_bytes=_bytes(doc), edits={2: "₹7,50,000"}
    )

    assert changed == 1
    cells = Document(io.BytesIO(edited)).tables[0].rows[0].cells
    assert cells[0].text == "Basic"
    assert cells[1].text == "₹7,50,000"


def test_editing_several_lines_at_once():
    data = _docx("One", "Two", "Three")
    edited, changed = apply_paragraph_edits(
        docx_bytes=data, edits={0: "Uno", 2: "Tres"}
    )
    assert changed == 2
    assert _texts(edited) == ["Uno", "Two", "Tres"]


def test_a_line_set_back_to_its_original_text_is_not_a_change():
    data = _docx("Unchanged line", "Other line")
    edited, changed = apply_paragraph_edits(
        docx_bytes=data, edits={0: "Unchanged line"}
    )
    assert changed == 0
    assert edited == data, "no change means the stored bytes are returned as-is"


def test_clearing_a_line_empties_it_without_removing_the_paragraph():
    data = _docx("Keep", "Delete this clause", "Keep")
    edited, changed = apply_paragraph_edits(docx_bytes=data, edits={1: ""})

    assert changed == 1
    assert _texts(edited) == ["Keep", "Keep"]
    assert len(Document(io.BytesIO(edited)).paragraphs) == 3


def test_unicode_survives_the_round_trip():
    data = _docx("Salary: TBD")
    edited, _ = apply_paragraph_edits(
        docx_bytes=data, edits={0: "Salary: ₹12,00,000 (बारह लाख)"}
    )
    assert _texts(edited) == ["Salary: ₹12,00,000 (बारह लाख)"]


def test_an_index_past_the_end_is_rejected_rather_than_guessed():
    data = _docx("Only line")
    with pytest.raises(TextEditError) as exc:
        apply_paragraph_edits(docx_bytes=data, edits={7: "Somewhere else"})
    assert "Reopen the editor" in str(exc.value)


def test_a_negative_index_is_rejected():
    data = _docx("Only line")
    with pytest.raises(TextEditError):
        apply_paragraph_edits(docx_bytes=data, edits={-1: "Nope"})


# ── Fingerprint ───────────────────────────────────────────────────────────


def test_fingerprint_is_stable_for_the_same_bytes():
    data = _docx("Dear Rahul,")
    assert draft_fingerprint(data) == draft_fingerprint(data)


def test_fingerprint_changes_once_the_draft_is_edited():
    data = _docx("Dear Rahul,")
    edited, _ = apply_paragraph_edits(docx_bytes=data, edits={0: "Dear Priya,"})
    assert draft_fingerprint(edited) != draft_fingerprint(data)
