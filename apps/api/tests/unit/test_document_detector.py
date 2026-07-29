"""Detection tests for the document pipeline.

The headline case is `test_filled_offer_letter_*`: a *completed* offer letter
used as the template. The previous `blank_detector` found zero fill-points in
such a document, because all three of its heuristics required the value to be
absent. That is the failure this detector exists to fix, so it is tested first
and in detail.
"""
from __future__ import annotations

import io

import pytest
from docx import Document

from app.services.documents.analysis.detector import (
    KIND_BLANK_MARKER,
    KIND_EMPTY_CELL,
    KIND_LABEL_COLON_EMPTY,
    KIND_LABEL_COLON_VALUE,
    KIND_PATTERN_VALUE,
    KIND_TABLE_CELL_VALUE,
    find_all_candidates,
    scan_for_unfilled_signals,
)
from app.services.documents.constants import (
    ACTION_INSERT_AFTER_LABEL,
    ACTION_INSERT_EMPTY_CELL,
    ACTION_REPLACE_SPAN,
)


def _docx(paragraphs: list[str], tables: list[list[tuple[str, str]]] | None = None) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    for rows in tables or []:
        table = doc.add_table(rows=len(rows), cols=2)
        for i, (left, right) in enumerate(rows):
            table.rows[i].cells[0].text = left
            table.rows[i].cells[1].text = right
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _by_text(candidates, text: str):
    return [c for c in candidates if c.original_text == text]


# ── The case the old detector could not see ───────────────────────────────


FILLED_OFFER_LETTER = [
    "Dear Rahul,",
    "We are pleased to offer you the position of Software Engineer.",
    "Joining Date: 15 August 2026",
    "Annual Salary: ₹12,00,000",
    "Reporting Manager: John Smith",
    "Office Location: Bengaluru",
]


def test_filled_offer_letter_detects_labelled_values():
    cands = find_all_candidates(_docx(FILLED_OFFER_LETTER))
    found = {c.original_text for c in cands}

    assert "15 August 2026" in found
    assert "₹12,00,000" in found
    assert "John Smith" in found
    assert "Bengaluru" in found


def test_filled_offer_letter_carries_the_label_as_a_hint():
    cands = find_all_candidates(_docx(FILLED_OFFER_LETTER))
    hints = {c.original_text: c.label_hint for c in cands}

    assert hints["15 August 2026"] == "Joining Date"
    assert hints["John Smith"] == "Reporting Manager"
    assert hints["₹12,00,000"] == "Annual Salary"


def test_filled_offer_letter_marks_values_as_occupied():
    """These spans hold another candidate's real data. HR must be able to tell
    them apart from blanks, because rejecting one leaves the sample person's
    details in a document sent to someone else."""
    cands = find_all_candidates(_docx(FILLED_OFFER_LETTER))
    for c in cands:
        if c.original_text in {"15 August 2026", "John Smith", "₹12,00,000"}:
            assert c.is_occupied, c


def test_label_value_beats_pattern_for_the_same_span():
    """`Joining Date: 15 August 2026` matches BOTH the label-value heuristic and
    the standalone date pattern. Exactly one candidate must survive, or the
    renderer would write the value twice into the same characters."""
    cands = find_all_candidates(_docx(["Joining Date: 15 August 2026"]))
    dates = _by_text(cands, "15 August 2026")
    assert len(dates) == 1
    assert dates[0].detection_kind == KIND_LABEL_COLON_VALUE


def test_offsets_point_at_the_real_characters():
    cands = find_all_candidates(_docx(["Joining Date: 15 August 2026"]))
    c = _by_text(cands, "15 August 2026")[0]
    line = "Joining Date: 15 August 2026"
    assert line[c.start_offset:c.end_offset] == "15 August 2026"


# ── Empty fill-points: the previous behaviour, still working ──────────────


def test_underscore_blank_detected():
    cands = find_all_candidates(_docx(["Candidate Name: ____________"]))
    assert any(c.detection_kind == KIND_BLANK_MARKER for c in cands)
    assert all(c.action == ACTION_REPLACE_SPAN for c in cands)


def test_bracket_placeholder_detected():
    cands = find_all_candidates(_docx(["Dear [CANDIDATE NAME],"]))
    marker = _by_text(cands, "[CANDIDATE NAME]")
    assert len(marker) == 1
    assert marker[0].detection_kind == KIND_BLANK_MARKER


def test_label_with_nothing_after_it_detected():
    cands = find_all_candidates(_docx(["Employee Signature:"]))
    assert len(cands) == 1
    assert cands[0].detection_kind == KIND_LABEL_COLON_EMPTY
    assert cands[0].action == ACTION_INSERT_AFTER_LABEL
    assert cands[0].start_offset == cands[0].end_offset


def test_empty_table_cell_beside_a_label_detected():
    cands = find_all_candidates(_docx([], tables=[[("Designation", "")]]))
    assert len(cands) == 1
    assert cands[0].detection_kind == KIND_EMPTY_CELL
    assert cands[0].action == ACTION_INSERT_EMPTY_CELL
    assert cands[0].label_hint == "Designation"


def test_filled_table_cell_beside_a_label_detected():
    cands = find_all_candidates(_docx([], tables=[[("Designation", "Software Engineer")]]))
    values = _by_text(cands, "Software Engineer")
    assert len(values) == 1
    assert values[0].detection_kind == KIND_TABLE_CELL_VALUE
    assert values[0].label_hint == "Designation"
    assert values[0].is_occupied


# ── Precision: what must NOT be detected ──────────────────────────────────


@pytest.mark.parametrize(
    "prose",
    [
        "The terms of your employment are as follows:",
        "This agreement includes the following:",
        "IN WITNESS WHEREOF:",
        "You will be governed by the policies of the company.",
    ],
)
def test_prose_lead_ins_are_not_field_labels(prose: str):
    cands = find_all_candidates(_docx([prose]))
    assert not [c for c in cands if c.detection_kind == KIND_LABEL_COLON_EMPTY]


def test_long_clause_after_a_colon_is_not_a_field_value():
    line = (
        "Confidentiality: You agree to keep all proprietary information "
        "strictly confidential for the duration of your employment and "
        "thereafter without limitation in time."
    )
    cands = find_all_candidates(_docx([line]))
    assert not [c for c in cands if c.detection_kind == KIND_LABEL_COLON_VALUE]


def test_no_candidates_in_a_plain_prose_document():
    cands = find_all_candidates(
        _docx([
            "This document sets out the general policies of the organisation.",
            "Employees are expected to conduct themselves professionally.",
        ])
    )
    assert cands == []


# ── Standalone patterns ───────────────────────────────────────────────────


def test_amount_inside_a_clause_detected():
    line = "You will receive a gross salary of ₹12,00,000 per annum."
    cands = find_all_candidates(_docx([line]))
    amounts = [c for c in cands if c.detection_kind == KIND_PATTERN_VALUE]
    assert amounts
    assert line[amounts[0].start_offset:amounts[0].end_offset].startswith("₹12,00,000")


def test_email_detected():
    cands = find_all_candidates(_docx(["Write to us at hr@contexo.com for questions."]))
    assert _by_text(cands, "hr@contexo.com")


def test_include_occupied_false_suppresses_value_candidates():
    cands = find_all_candidates(_docx(FILLED_OFFER_LETTER), include_occupied=False)
    assert cands == []


# ── Structural invariants ─────────────────────────────────────────────────


def test_candidate_ids_are_dense_and_in_document_order():
    cands = find_all_candidates(_docx(FILLED_OFFER_LETTER))
    assert [c.candidate_id for c in cands] == list(range(len(cands)))
    positions = [(c.paragraph_index, c.start_offset) for c in cands]
    assert positions == sorted(positions)


def test_no_two_candidates_overlap_in_the_same_paragraph():
    cands = find_all_candidates(
        _docx([
            "Joining Date: 15 August 2026",
            "Annual Salary: ₹12,00,000",
            "Contact: hr@contexo.com or +91 98765 43210",
        ])
    )
    by_para: dict[int, list[tuple[int, int]]] = {}
    for c in cands:
        by_para.setdefault(c.paragraph_index, []).append((c.start_offset, c.end_offset))
    for spans in by_para.values():
        spans.sort()
        for (s1, e1), (s2, e2) in zip(spans, spans[1:]):
            assert e1 <= s2, f"overlapping spans {(s1, e1)} and {(s2, e2)}"


def test_candidate_count_is_bounded_on_a_large_document():
    cands = find_all_candidates(_docx([f"Field {i}: value {i}" for i in range(600)]))
    assert len(cands) <= 600  # sanity: no combinatorial blowup


# ── Post-render safety net ────────────────────────────────────────────────


def test_scan_flags_leftover_blanks():
    warnings = scan_for_unfilled_signals("Name: ________\nSignature:")
    assert len(warnings) == 2


def test_scan_is_clean_on_a_fully_filled_document():
    assert scan_for_unfilled_signals("Name: Rahul\nJoining Date: 15 August 2026") == []
