"""Deciding whether a document really is a hand-authored Jinja template.

The regression under test is a specific production failure: `{{ Signing Date }}`
matched the old detection regex (whose character class included `\\s`), so the
document was classified "already templated", skipped blank detection, and died
inside docxtpl at render time with:

    expected token 'end of print statement', got 'signing'

HR saw that with no indication of which placeholder was wrong. The rule these
tests pin down is that a tag is only trusted if Jinja's own parser accepts it.
"""
from __future__ import annotations

import io

import pytest
from docx import Document

from app.services.agents.onboarding_v2.jinja_validator import (
    STRATEGY_JINJA,
    STRATEGY_MIXED_INVALID,
    STRATEGY_SLOTS,
    classify_template,
    extract_jinja_tags,
    validate_jinja_tags,
)


def _docx(*paragraphs: str) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_no_tags_means_the_document_needs_detection() -> None:
    assert classify_template(_docx("Dear candidate,")).strategy == STRATEGY_SLOTS


def test_valid_tags_are_honoured_as_a_template() -> None:
    scan = classify_template(_docx("Dear {{ candidate_name }},"))
    assert scan.strategy == STRATEGY_JINJA
    assert scan.errors == []


def test_filters_and_dotted_access_stay_valid() -> None:
    """The old permissive regex existed to allow these; the parser handles them
    natively, so tightening validation must not reject them."""
    errors, _referenced = validate_jinja_tags(
        ["{{ ctc | upper }}", "{{ ctc_breakdown.basic }}"]
    )
    assert errors == []


def test_malformed_multiword_tag_is_never_trusted() -> None:
    """THE regression. A tag with a space in the name cannot parse, so the
    document must NOT be routed to docxtpl."""
    scan = classify_template(_docx("Signed on {{ Signing Date }} at our office."))
    assert scan.strategy == STRATEGY_MIXED_INVALID
    assert scan.strategy != STRATEGY_JINJA
    assert len(scan.errors) == 1


def test_malformed_tag_error_is_actionable() -> None:
    """HR must be able to fix it from the message alone — the old failure gave
    them a parser token name and nothing else."""
    scan = classify_template(_docx("Signed on {{ Signing Date }}."))
    message = scan.errors[0]
    assert "{{ Signing Date }}" in message
    assert "cannot contain spaces" in message
    assert "{{ signing_date }}" in message  # the suggested fix


def test_a_document_mixing_valid_and_invalid_tags_is_not_trusted() -> None:
    scan = classify_template(
        _docx("Dear {{ candidate_name }}, signed {{ Signing Date }}.")
    )
    assert scan.strategy == STRATEGY_MIXED_INVALID


def test_valid_tag_naming_an_unsupplied_variable_is_reported_early() -> None:
    """Still a real template, so it renders via docxtpl — but the name would
    raise TemplateVariableError mid-generation, so it's surfaced at analyze
    time instead."""
    scan = classify_template(_docx("PF number: {{ employee_pf_number }}"))
    assert scan.strategy == STRATEGY_JINJA
    assert scan.unknown_variables == ["employee_pf_number"]


def test_tags_are_found_inside_tables_and_headers() -> None:
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "{{ company_name }}"
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "{{ candidate_name }}"
    buf = io.BytesIO()
    doc.save(buf)
    tags = extract_jinja_tags(buf.getvalue())
    assert "{{ company_name }}" in tags
    assert "{{ candidate_name }}" in tags


@pytest.mark.parametrize(
    "tag",
    ["{{ }}", "{{ 1name }}", "{{ candidate name }}", "{{ a b c }}"],
)
def test_assorted_malformed_tags_are_rejected(tag: str) -> None:
    errors, _ = validate_jinja_tags([tag])
    assert errors, f"{tag} should not have parsed"
