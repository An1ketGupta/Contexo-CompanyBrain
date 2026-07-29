"""Template library and builder endpoints.

Covers the first half of the pipeline's API flow:

    Upload Template → Analyze → Confirm Variables → Save Schema

Generation, preview-for-a-candidate, approval and sending live in the
generation router.

Every handler resolves `org_id` from the JWT and passes it into the service
layer, which filters on it explicitly — the service-role client bypasses RLS, so
the org check is code, not policy, on the write path.
"""
from __future__ import annotations

import asyncio
import io
import logging

from docx import Document
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status

from app.auth import verify_jwt
from app.database import get_service_client
from app.errors import NoOrganization
from app.models.documents_pipeline import (
    AnalyzeResponse,
    DocumentTypeCreate,
    DocumentTypeRead,
    PreviewRequest,
    PreviewResponse,
    SchemaRead,
    SlotCreate,
    SlotRead,
    SlotUpdate,
    TemplateParagraph,
    TemplateRead,
    TemplateReadiness,
    TemplateTextResponse,
    TemplateUpdate,
    TemplateVersionRead,
    VariableCreate,
    VariableRead,
    VariableUpdate,
)
from app.services.documents import schema as schema_service
from app.services.documents import storage, templates
from app.services.documents.analysis.detector import scan_for_unfilled_signals
from app.services.documents.analysis.runner import AnalysisNotSupported, run_analysis
from app.services.documents.constants import (
    DEFAULT_CONFIRM_THRESHOLD,
    FORMAT_DOCX,
    FORMAT_PDF,
    MIME_DOCX,
    MIME_PDF,
    UPLOADABLE_MIMES,
)
from app.services.documents.docx_positions import (
    canonical_paragraphs,
    paragraph_hash,
    paragraph_run_text,
)
from app.services.documents.generation.renderer import (
    DocumentRenderError,
    render_to_pdf,
)
from app.services.documents.preview import preview_signature_labels, sample_values
from app.services.documents.validation.engine import build_context

log = logging.getLogger(__name__)

router = APIRouter(prefix="/document-templates", tags=["document-templates"])

# 25 MB. Word documents with embedded logos and letterheads run large, but a
# template an order of magnitude past this is a mis-upload, and parsing it would
# tie up a worker.
MAX_UPLOAD_BYTES = 25 * 1024 * 1024


def _require_org(current_user: dict) -> tuple[str, str]:
    org_id = current_user.get("org_id")
    user_id = current_user.get("user_id")
    if not org_id or not user_id:
        raise NoOrganization("No organization found. Please sign out and sign back in.")
    return org_id, user_id


def _template_to_read(row: dict) -> TemplateRead:
    joined = row.get("document_types") or {}
    return TemplateRead(
        id=row["id"],
        name=row["name"],
        description=row.get("description"),
        status=row["status"],
        is_default=row.get("is_default", False),
        document_type_id=row["document_type_id"],
        document_type_key=joined.get("key"),
        document_type_label=joined.get("label"),
        current_version_id=row.get("current_version_id"),
        created_at=row.get("created_at"),
        updated_at=row.get("updated_at"),
    )


async def _read_upload(file: UploadFile) -> tuple[bytes, str]:
    """Read and validate an uploaded template file."""
    data = await file.read()
    if not data:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "The uploaded file is empty.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            f"Templates must be under {MAX_UPLOAD_BYTES // (1024 * 1024)} MB.",
        )

    mime = (file.content_type or "").split(";")[0].strip()
    if mime not in UPLOADABLE_MIMES:
        # Browsers are inconsistent about DOCX content types, so fall back to
        # the extension rather than rejecting a legitimate upload.
        name = (file.filename or "").lower()
        if name.endswith(".docx"):
            mime = MIME_DOCX
        elif name.endswith(".pdf"):
            mime = MIME_PDF
        else:
            raise HTTPException(
                status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                "Upload a Word (.docx) or PDF file.",
            )
    return data, mime


# ── Document types ─────────────────────────────────────────────────────────


@router.get("/types", response_model=list[DocumentTypeRead])
async def list_types(current_user: dict = Depends(verify_jwt)):
    org_id, _ = _require_org(current_user)
    rows = await templates.list_document_types(org_id)
    return [DocumentTypeRead(**{k: r.get(k) for k in DocumentTypeRead.model_fields}) for r in rows]


@router.post("/types", response_model=DocumentTypeRead, status_code=status.HTTP_201_CREATED)
async def create_type(body: DocumentTypeCreate, current_user: dict = Depends(verify_jwt)):
    """Add a document type for this org — no migration, no deploy."""
    org_id, _ = _require_org(current_user)
    try:
        row = await templates.create_document_type(
            org_id=org_id, key=body.key, label=body.label, description=body.description
        )
    except templates.TemplateError as exc:
        raise HTTPException(status.HTTP_409_CONFLICT, str(exc)) from exc
    return DocumentTypeRead(**{k: row.get(k) for k in DocumentTypeRead.model_fields})


@router.get("/readiness", response_model=list[TemplateReadiness])
async def readiness(
    keys: str = "letter_of_intent,appointment_letter,nda,induction",
    current_user: dict = Depends(verify_jwt),
):
    """Can each of these document types actually be generated right now?

    Two distinct not-ready states, which the previous "is a template tagged?"
    check could not tell apart: no default template at all, and a template whose
    fields nobody has confirmed. The second one blocks generation just as hard,
    and used to surface only when a run failed.
    """
    org_id, _ = _require_org(current_user)
    wanted = [k.strip() for k in keys.split(",") if k.strip()]
    types = {t["key"]: t for t in await templates.list_document_types(org_id)}

    out: list[TemplateReadiness] = []
    for key in wanted:
        doc_type = types.get(key)
        label = doc_type["label"] if doc_type else key.replace("_", " ").title()

        resolved = await templates.resolve_default_version(org_id=org_id, type_key=key)
        if not resolved:
            out.append(TemplateReadiness(
                type_key=key, label=label, ready=False, reason="no_template"
            ))
            continue

        template, version = resolved
        _variables, slots = await schema_service.confirmed_schema(
            org_id=org_id, version_id=version["id"]
        )
        out.append(TemplateReadiness(
            type_key=key,
            label=label,
            ready=bool(slots),
            reason=None if slots else "no_confirmed_fields",
            template_id=template["id"],
            template_name=template["name"],
            version_id=version["id"],
            confirmed_field_count=len(slots),
        ))
    return out


# ── Templates ──────────────────────────────────────────────────────────────


@router.get("", response_model=list[TemplateRead])
async def list_templates(
    include_archived: bool = False, current_user: dict = Depends(verify_jwt)
):
    org_id, _ = _require_org(current_user)
    rows = await templates.list_templates(org_id=org_id, include_archived=include_archived)
    return [_template_to_read(r) for r in rows]


@router.post("", response_model=TemplateRead, status_code=status.HTTP_201_CREATED)
async def create_template(
    name: str = Form(...),
    document_type_id: str = Form(...),
    description: str | None = Form(None),
    file: UploadFile = File(...),
    current_user: dict = Depends(verify_jwt),
):
    """Create a template and store its first version in one call.

    A template with no file is not a useful thing to have in the library, so the
    two steps are one operation.
    """
    org_id, user_id = _require_org(current_user)
    data, mime = await _read_upload(file)

    template = await templates.create_template(
        org_id=org_id,
        document_type_id=document_type_id,
        name=name,
        description=description,
        created_by=user_id,
    )
    try:
        await templates.add_version(
            org_id=org_id,
            template_id=template["id"],
            data=data,
            original_filename=file.filename or "template.docx",
            mime_type=mime,
            uploaded_by=user_id,
        )
    except storage.TemplateStorageError as exc:
        raise HTTPException(status.HTTP_502_BAD_GATEWAY, str(exc)) from exc

    return _template_to_read(await templates.get_template(org_id, template["id"]))


@router.get("/{template_id}", response_model=TemplateRead)
async def get_template(template_id: str, current_user: dict = Depends(verify_jwt)):
    org_id, _ = _require_org(current_user)
    try:
        return _template_to_read(await templates.get_template(org_id, template_id))
    except templates.TemplateNotFound as exc:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Template not found.") from exc


@router.patch("/{template_id}", response_model=TemplateRead)
async def update_template(
    template_id: str, body: TemplateUpdate, current_user: dict = Depends(verify_jwt)
):
    org_id, user_id = _require_org(current_user)
    try:
        await templates.update_template(
            org_id=org_id,
            template_id=template_id,
            patch=body.model_dump(exclude_none=True),
            actor_user_id=user_id,
        )
        return _template_to_read(await templates.get_template(org_id, template_id))
    except templates.TemplateNotFound as exc:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Template not found.") from exc


@router.post("/{template_id}/default", response_model=TemplateRead)
async def make_default(template_id: str, current_user: dict = Depends(verify_jwt)):
    org_id, user_id = _require_org(current_user)
    try:
        await templates.set_default(
            org_id=org_id, template_id=template_id, actor_user_id=user_id
        )
        return _template_to_read(await templates.get_template(org_id, template_id))
    except templates.TemplateNotFound as exc:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Template not found.") from exc


@router.post("/{template_id}/archive", response_model=TemplateRead)
async def archive_template(template_id: str, current_user: dict = Depends(verify_jwt)):
    org_id, user_id = _require_org(current_user)
    try:
        await templates.archive_template(
            org_id=org_id, template_id=template_id, actor_user_id=user_id
        )
        return _template_to_read(await templates.get_template(org_id, template_id))
    except templates.TemplateNotFound as exc:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Template not found.") from exc


# ── Versions ───────────────────────────────────────────────────────────────


@router.get("/{template_id}/versions", response_model=list[TemplateVersionRead])
async def list_versions(template_id: str, current_user: dict = Depends(verify_jwt)):
    org_id, _ = _require_org(current_user)
    rows = await templates.list_versions(org_id=org_id, template_id=template_id)
    return [
        TemplateVersionRead(**{k: r.get(k) for k in TemplateVersionRead.model_fields})
        for r in rows
    ]


@router.post(
    "/{template_id}/versions",
    response_model=TemplateVersionRead,
    status_code=status.HTTP_201_CREATED,
)
async def upload_version(
    template_id: str,
    file: UploadFile = File(...),
    current_user: dict = Depends(verify_jwt),
):
    """Add a new version. The previous file is left exactly where it is."""
    org_id, user_id = _require_org(current_user)
    data, mime = await _read_upload(file)
    try:
        row = await templates.add_version(
            org_id=org_id,
            template_id=template_id,
            data=data,
            original_filename=file.filename or "template.docx",
            mime_type=mime,
            uploaded_by=user_id,
        )
    except templates.TemplateNotFound as exc:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Template not found.") from exc
    except storage.TemplateStorageError as exc:
        raise HTTPException(status.HTTP_502_BAD_GATEWAY, str(exc)) from exc
    return TemplateVersionRead(**{k: row.get(k) for k in TemplateVersionRead.model_fields})


@router.get("/versions/{version_id}/download")
async def download_version(version_id: str, current_user: dict = Depends(verify_jwt)):
    """Signed URL for the original, unmodified file."""
    org_id, _ = _require_org(current_user)
    try:
        version = await templates.get_version(org_id, version_id)
    except templates.TemplateNotFound as exc:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Version not found.") from exc

    url = await storage.mint_signed_url(version["storage_path"])
    if not url:
        raise HTTPException(
            status.HTTP_502_BAD_GATEWAY, "Couldn't produce a download link. Try again."
        )
    return {"url": url, "filename": version["original_filename"]}


# ── Analysis ───────────────────────────────────────────────────────────────


@router.post("/versions/{version_id}/analyze", response_model=AnalyzeResponse)
async def analyze_version(version_id: str, current_user: dict = Depends(verify_jwt)):
    """Detect fields and propose names and types.

    Safe to re-run: a variable or slot HR has already confirmed or rejected is
    never touched.
    """
    org_id, user_id = _require_org(current_user)
    try:
        result = await run_analysis(
            org_id=org_id, version_id=version_id, actor_user_id=user_id
        )
    except templates.TemplateNotFound as exc:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Version not found.") from exc
    except AnalysisNotSupported as exc:
        raise HTTPException(status.HTTP_422_UNPROCESSABLE_ENTITY, str(exc)) from exc
    except storage.TemplateStorageError as exc:
        raise HTTPException(status.HTTP_502_BAD_GATEWAY, str(exc)) from exc
    return AnalyzeResponse(**result)


# ── Schema ─────────────────────────────────────────────────────────────────


@router.get("/versions/{version_id}/schema", response_model=SchemaRead)
async def get_schema(version_id: str, current_user: dict = Depends(verify_jwt)):
    """Everything the template builder renders: the version, its variables, and
    its slots."""
    org_id, _ = _require_org(current_user)
    try:
        version = await templates.get_version(org_id, version_id)
    except templates.TemplateNotFound as exc:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Version not found.") from exc

    variables = await schema_service.list_variables(org_id=org_id, version_id=version_id)
    slots = await schema_service.list_slots(org_id=org_id, version_id=version_id)

    return SchemaRead(
        version=TemplateVersionRead(
            **{k: version.get(k) for k in TemplateVersionRead.model_fields}
        ),
        variables=[
            VariableRead(**{k: v.get(k) for k in VariableRead.model_fields})
            for v in variables
        ],
        slots=[
            SlotRead(**{k: s.get(k) for k in SlotRead.model_fields}) for s in slots
        ],
        confirm_threshold=DEFAULT_CONFIRM_THRESHOLD,
    )


@router.post(
    "/versions/{version_id}/variables",
    response_model=VariableRead,
    status_code=status.HTTP_201_CREATED,
)
async def create_variable(
    version_id: str, body: VariableCreate, current_user: dict = Depends(verify_jwt)
):
    org_id, user_id = _require_org(current_user)
    try:
        await templates.get_version(org_id, version_id)
    except templates.TemplateNotFound as exc:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Version not found.") from exc

    try:
        row = await schema_service.create_variable(
            org_id=org_id,
            version_id=version_id,
            actor_user_id=user_id,
            **body.model_dump(),
        )
    except Exception as exc:  # noqa: BLE001 — surfaced as a readable conflict
        raise HTTPException(
            status.HTTP_409_CONFLICT,
            f"Couldn't add '{body.internal_name}'. A field with that name may "
            "already exist on this version.",
        ) from exc
    return VariableRead(**{k: row.get(k) for k in VariableRead.model_fields})


@router.patch("/variables/{variable_id}", response_model=VariableRead)
async def update_variable(
    variable_id: str, body: VariableUpdate, current_user: dict = Depends(verify_jwt)
):
    org_id, user_id = _require_org(current_user)
    try:
        row = await schema_service.update_variable(
            org_id=org_id,
            variable_id=variable_id,
            patch=body.model_dump(exclude_none=True),
            actor_user_id=user_id,
        )
    except schema_service.SchemaNotFound as exc:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Field not found.") from exc
    return VariableRead(**{k: row.get(k) for k in VariableRead.model_fields})


@router.delete("/variables/{variable_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_variable(variable_id: str, current_user: dict = Depends(verify_jwt)):
    org_id, _ = _require_org(current_user)
    await schema_service.delete_variable(org_id=org_id, variable_id=variable_id)


@router.patch("/slots/{slot_id}", response_model=SlotRead)
async def update_slot(
    slot_id: str, body: SlotUpdate, current_user: dict = Depends(verify_jwt)
):
    org_id, user_id = _require_org(current_user)
    try:
        await schema_service.update_slot(
            org_id=org_id,
            slot_id=slot_id,
            patch=body.model_dump(exclude_unset=True),
            actor_user_id=user_id,
        )
    except schema_service.SchemaNotFound as exc:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Field position not found.") from exc

    rows = await schema_service.list_slots(
        org_id=org_id,
        version_id=(await _slot_version(org_id, slot_id)),
    )
    match = next((r for r in rows if r["id"] == slot_id), None)
    if not match:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Field position not found.")
    return SlotRead(**{k: match.get(k) for k in SlotRead.model_fields})


async def _slot_version(org_id: str, slot_id: str) -> str:
    def _fetch() -> str | None:
        res = (
            get_service_client().table("doc_template_slots")
            .select("version_id")
            .eq("id", slot_id)
            .eq("org_id", org_id)
            .limit(1)
            .execute()
        )
        rows = res.data or []
        return rows[0]["version_id"] if rows else None

    version_id = await asyncio.to_thread(_fetch)
    if not version_id:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Field position not found.")
    return version_id


@router.get("/versions/{version_id}/text", response_model=TemplateTextResponse)
async def get_version_text(version_id: str, current_user: dict = Depends(verify_jwt)):
    """The template's paragraphs, addressed the way slots are.

    Backs the "point this field at a place in the document" picker. Read-only:
    HR selects a span, and the offsets they pick index into exactly the same
    string the renderer will splice into. Empty paragraphs are dropped from the
    response but keep their true index, so nothing shifts.
    """
    org_id, _ = _require_org(current_user)
    try:
        data, version = await templates.load_version_bytes(
            org_id=org_id, version_id=version_id
        )
    except templates.TemplateNotFound as exc:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Version not found.") from exc
    except storage.TemplateStorageError as exc:
        raise HTTPException(status.HTTP_502_BAD_GATEWAY, str(exc)) from exc

    if version.get("mime_type") != MIME_DOCX:
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY,
            "Only Word (.docx) templates can be browsed this way.",
        )

    canonical = canonical_paragraphs(Document(io.BytesIO(data)))
    paragraphs = [
        TemplateParagraph(index=idx, kind=kind, text=text)
        for idx, (paragraph, kind) in enumerate(canonical)
        if (text := paragraph_run_text(paragraph)).strip()
    ]
    return TemplateTextResponse(paragraphs=paragraphs)


@router.post("/versions/{version_id}/preview", response_model=PreviewResponse)
async def preview_version(
    version_id: str,
    body: PreviewRequest | None = None,
    current_user: dict = Depends(verify_jwt),
):
    """Render this template with obviously-fake data.

    Answers the only question that matters before a template goes live: does the
    document still look right once the fields are filled? That is a question
    about layout, so it must be answerable without choosing a candidate.

    Uses the CONFIRMED schema, so what HR previews is exactly what a generation
    would produce.
    """
    org_id, _ = _require_org(current_user)
    try:
        data, version = await templates.load_version_bytes(
            org_id=org_id, version_id=version_id
        )
    except templates.TemplateNotFound as exc:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Version not found.") from exc
    except storage.TemplateStorageError as exc:
        raise HTTPException(status.HTTP_502_BAD_GATEWAY, str(exc)) from exc

    if version.get("mime_type") != MIME_DOCX:
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY,
            "Previews are only available for Word (.docx) templates.",
        )

    variables, slots = await schema_service.confirmed_schema(
        org_id=org_id, version_id=version_id
    )
    if not slots:
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY,
            "Confirm at least one field before previewing.",
        )

    values = sample_values(variables, overrides=(body.values if body else {}))
    context = build_context(variables=variables, values=values)
    # Show a readable label where a signature will be stamped, not the raw
    # sentinel the e-sign service searches for.
    context.update(preview_signature_labels(variables))

    try:
        filled_docx, pdf_bytes = await render_to_pdf(
            docx_bytes=data,
            slots=slots,
            context=context,
            template_name=version.get("original_filename"),
        )
    except DocumentRenderError as exc:
        raise HTTPException(status.HTTP_422_UNPROCESSABLE_ENTITY, str(exc)) from exc

    docx_path = storage.preview_path(
        org_id=org_id, version_id=version_id, fmt=FORMAT_DOCX
    )
    await storage.upload(
        path=docx_path, data=filled_docx, mime_type=MIME_DOCX, overwrite=True
    )
    docx_url = await storage.mint_signed_url(docx_path)

    pdf_url = None
    warnings: list[str] = []
    if pdf_bytes:
        pdf_path = storage.preview_path(
            org_id=org_id, version_id=version_id, fmt=FORMAT_PDF
        )
        await storage.upload(
            path=pdf_path, data=pdf_bytes, mime_type=MIME_PDF, overwrite=True
        )
        pdf_url = await storage.mint_signed_url(pdf_path)
    else:
        warnings.append(
            "PDF preview is unavailable right now — the Word preview still "
            "reflects exactly what will be generated."
        )

    warnings.extend(scan_for_unfilled_signals(_document_text(filled_docx)))

    return PreviewResponse(
        docx_url=docx_url, pdf_url=pdf_url, warnings=warnings, used_values=values
    )


def _document_text(docx_bytes: bytes) -> str:
    """Flat text of a rendered document, for the leftover-blank scan."""
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(
        paragraph_run_text(p) for p, _kind in canonical_paragraphs(doc)
    )


@router.post(
    "/versions/{version_id}/slots",
    response_model=SlotRead,
    status_code=status.HTTP_201_CREATED,
)
async def create_slot(
    version_id: str, body: SlotCreate, current_user: dict = Depends(verify_jwt)
):
    """Add a fill-point by hand, for a blank detection could not see.

    The paragraph hash is computed here against the stored file rather than
    trusted from the client, so a manually-added slot carries the same drift
    protection as a detected one.
    """
    org_id, user_id = _require_org(current_user)
    try:
        data, version = await templates.load_version_bytes(
            org_id=org_id, version_id=version_id
        )
    except templates.TemplateNotFound as exc:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Version not found.") from exc

    canonical = canonical_paragraphs(Document(io.BytesIO(data)))
    if not (0 <= body.paragraph_index < len(canonical)):
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY,
            "That position is not in this document.",
        )

    paragraph, kind = canonical[body.paragraph_index]
    text = paragraph_run_text(paragraph)
    if body.end_offset > len(text):
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY,
            "That position is past the end of the paragraph.",
        )

    try:
        row = await schema_service.create_slot(
            org_id=org_id,
            version_id=version_id,
            variable_id=body.variable_id,
            paragraph_index=body.paragraph_index,
            paragraph_kind=kind,
            paragraph_hash=paragraph_hash(text),
            source_file_sha256=version["file_sha256"],
            start_offset=body.start_offset,
            end_offset=body.end_offset,
            action=body.action,
            original_text=text[body.start_offset:body.end_offset],
            actor_user_id=user_id,
        )
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(
            status.HTTP_409_CONFLICT,
            "There is already a field at that position.",
        ) from exc

    return SlotRead(**{k: row.get(k) for k in SlotRead.model_fields})
