"""Format-preserving export.

Two outputs per RFP:

  1. FILLED ORIGINAL (xlsx → xlsx, docx → docx-with-inline-answers).
     Loaded from the buyer's source bytes; we mutate ONLY the cells/paragraphs
     we own, preserving formatting, formulas, branding, merged cells, etc.

  2. SUMMARY DOCX (always emitted).
     Clean vendor-branded response document with a gap callout up top and
     per-requirement Q/A sections, citing sources.

Why both:
  * The buyer wants their own template back filled in.
  * The vendor's legal/sales team wants a tidy doc for archival and follow-up
    that doesn't bury answers inside buyer formatting.
"""
from __future__ import annotations

import logging
import re
from datetime import UTC, datetime
from io import BytesIO
from typing import Any

log = logging.getLogger(__name__)


# ── Public entry ────────────────────────────────────────────────────────────


def export_filled_xlsx(
    *,
    source_bytes: bytes,
    answers_by_row: dict[tuple[str, int], dict[str, Any]],
    answer_column_by_sheet: dict[str, str],
    notes_column_by_sheet: dict[str, str | None] | None = None,
) -> bytes:
    """Write answers into the buyer's original XLSX. Returns new bytes.

    Args:
      source_bytes: the original buyer file
      answers_by_row: keyed by (sheet_title, row_index_1based). Each value:
          {answer_text, confidence_band, sources: [{document_name, ...}], is_gap}
      answer_column_by_sheet: {"Sheet1": "C", ...}
      notes_column_by_sheet: optional; if set, we write the sources into the
          notes column so the buyer can see attribution.
    """
    import openpyxl
    from openpyxl.utils import column_index_from_string
    from openpyxl.styles import PatternFill, Font

    notes_column_by_sheet = notes_column_by_sheet or {}
    wb = openpyxl.load_workbook(BytesIO(source_bytes))

    gap_fill = PatternFill(start_color="FFEBEE", end_color="FFEBEE", fill_type="solid")
    gap_font = Font(color="B71C1C", italic=True)

    try:
        for sheet in wb.worksheets:
            sheet_title = sheet.title
            answer_col_letter = answer_column_by_sheet.get(sheet_title)
            if not answer_col_letter:
                continue
            answer_col = column_index_from_string(answer_col_letter)
            notes_col_letter = notes_column_by_sheet.get(sheet_title)
            notes_col = (
                column_index_from_string(notes_col_letter) if notes_col_letter else None
            )

            for (sheet_key, row_idx), payload in answers_by_row.items():
                if sheet_key != sheet_title:
                    continue
                answer_text = (payload.get("answer_text") or "").strip()
                is_gap = bool(payload.get("is_gap"))
                if is_gap and not answer_text:
                    answer_text = "GAP — flagged for legal review."

                cell = sheet.cell(row=row_idx, column=answer_col)
                cell.value = answer_text
                cell.alignment = openpyxl.styles.Alignment(wrap_text=True, vertical="top")
                if is_gap:
                    cell.fill = gap_fill
                    cell.font = gap_font

                if notes_col is not None:
                    sources = payload.get("sources") or []
                    if sources:
                        attribution = "Sources: " + "; ".join(
                            s.get("document_name", "?") for s in sources[:3]
                        )
                        sheet.cell(row=row_idx, column=notes_col).value = attribution

        out = BytesIO()
        wb.save(out)
        return out.getvalue()
    finally:
        wb.close()


def export_filled_docx(
    *,
    source_bytes: bytes,
    answers_by_external_ref: dict[str, dict[str, Any]],
) -> bytes:
    """For DOCX-table RFPs: walk the doc tables, find the answer column,
    and write the answer next to its question. We can't be as surgical as
    XLSX (no cell-address breadcrumb), so we match by question text. Returns
    bytes.

    For docx_list RFPs the safest thing is to emit a fresh DOCX rather than
    try to interleave; callers should use export_summary_docx for those.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document(BytesIO(source_bytes))

    # Map by-question-text-substring so we tolerate minor whitespace diffs.
    norm_answer_map = {
        _norm(k): v for k, v in answers_by_external_ref.items() if k
    }

    for table in doc.tables:
        if not table.rows:
            continue
        headers = [c.text.strip() for c in table.rows[0].cells]
        q_col = _find_first_col(headers, ("question", "requirement", "criteria", "item"))
        a_col = _find_first_col(headers, ("answer", "response", "vendor response", "supplier response"))
        if q_col is None:
            continue
        if a_col is None:
            # Append an Answer column conceptually; in DOCX we just write into the last cell.
            a_col = len(headers) - 1

        for row in table.rows[1:]:
            cells = row.cells
            if q_col >= len(cells) or a_col >= len(cells):
                continue
            q_text = cells[q_col].text.strip()
            if not q_text:
                continue
            entry = norm_answer_map.get(_norm(q_text))
            if not entry:
                continue
            answer_text = (entry.get("answer_text") or "").strip()
            is_gap = bool(entry.get("is_gap"))
            cell = cells[a_col]
            # Clear cell, then write our paragraph.
            for p in list(cell.paragraphs):
                p._element.getparent().remove(p._element)
            new_p = cell.add_paragraph()
            run = new_p.add_run(answer_text or ("GAP — flagged for legal review." if is_gap else ""))
            run.font.size = Pt(10)
            if is_gap:
                run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
                run.italic = True
            sources = entry.get("sources") or []
            if sources and not is_gap:
                src_p = cell.add_paragraph()
                src_run = src_p.add_run(
                    "Source: " + ", ".join(s.get("document_name", "?") for s in sources[:2])
                )
                src_run.italic = True
                src_run.font.size = Pt(8)
                src_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def export_summary_docx(
    *,
    title: str,
    requirements: list[dict[str, Any]],
    answers: dict[str, dict[str, Any]],
    company_name: str | None = None,
) -> bytes:
    """Generate the clean summary DOCX. Always emitted regardless of source
    format. Gap callout up top, per-requirement Q/A below.

    Args:
      requirements: list of dicts with {id, ordinal, external_ref, requirement_text, category}
      answers: keyed by requirement_id; values like {answer_text, sources, is_gap, flag_message, confidence_band}
    """
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt, RGBColor

    doc = Document()

    heading = doc.add_heading("RFP Response", level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub = doc.add_paragraph(title)
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if company_name:
        co = doc.add_paragraph(f"Prepared by: {company_name}")
        co.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    ts = doc.add_paragraph(
        f"Generated: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    ts.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()  # spacer

    # Gap callout
    gap_ids = [r["id"] for r in requirements if (answers.get(r["id"]) or {}).get("is_gap")]
    if gap_ids:
        warn = doc.add_heading(
            f"⚠ {len(gap_ids)} requirement(s) flagged for legal/SME review",
            level=1,
        )
        for run in warn.runs:
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        for r in requirements:
            ans = answers.get(r["id"]) or {}
            if not ans.get("is_gap"):
                continue
            p = doc.add_paragraph()
            p.add_run(f"  • [{r.get('external_ref') or r.get('ordinal') or ''}] ").bold = True
            p.add_run((r["requirement_text"] or "")[:240])
        doc.add_page_break()

    # Per-requirement
    for r in requirements:
        label = r.get("external_ref") or f"R{r.get('ordinal', 0) + 1}"
        doc.add_heading(f"{label}. {r.get('requirement_text', '').strip()[:200]}", level=2)

        if r.get("category"):
            cat = doc.add_paragraph()
            cr = cat.add_run(f"Category: {r['category']}")
            cr.italic = True
            cr.font.size = Pt(9)
            cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        ans = answers.get(r["id"]) or {}
        if ans.get("is_gap"):
            gap_p = doc.add_paragraph()
            gap_run = gap_p.add_run(
                "⚠ GAP — flagged for legal review. " + (ans.get("flag_message") or "")
            )
            gap_run.bold = True
            gap_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            if ans.get("answer_text"):
                ed = doc.add_paragraph(ans["answer_text"])
                for run in ed.runs:
                    run.italic = True
            continue

        body_text = (ans.get("answer_text") or "").strip()
        if body_text:
            doc.add_paragraph(body_text)
        else:
            empty = doc.add_paragraph("(No answer generated.)")
            for run in empty.runs:
                run.italic = True

        sources = ans.get("sources") or []
        if sources:
            srcs = doc.add_paragraph()
            srcs_run = srcs.add_run(
                "Sources: " + ", ".join(s.get("document_name", "?") for s in sources[:3])
            )
            srcs_run.italic = True
            srcs_run.font.size = Pt(9)
            srcs_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ── Internals ──────────────────────────────────────────────────────────────


_WS_RE = re.compile(r"\s+")


def _norm(s: str) -> str:
    return _WS_RE.sub(" ", (s or "").strip().lower())


def _find_first_col(headers: list[str], needles: tuple[str, ...]) -> int | None:
    norm = [h.strip().lower() for h in headers]
    for i, h in enumerate(norm):
        if any(n == h or n in h for n in needles):
            return i
    return None
