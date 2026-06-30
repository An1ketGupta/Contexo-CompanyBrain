"""Table-aware RFP parser.

Why a bespoke parser (instead of reusing ingestion.parser):
  * Ingestion's XLSX parser flattens each row into prose ("col: value, …") —
    great for retrieval, useless for round-tripping. We need to KEEP the row
    index and the answer column letter so the exporter can write the answer
    back into the same cell of the same file.
  * RFPs follow predictable shapes (Q/A/Notes columns; numbered DOCX lists;
    DOCX tables). Detecting them explicitly is far more reliable than asking
    Gemini to "find the questions in this prose blob."
  * The output is the agent's structured requirement list — no LLM needed
    for typical XLSX/DOCX RFPs. PDFs and free-text fall back to the LLM
    extractor downstream.

Public surface:

    parse_rfp(file_bytes, source_format) -> ParsedRfp

Where `ParsedRfp.extraction_source` is one of:
  * "xlsx_table"     — happy path: detected Q + A columns, fillable on export
  * "docx_table"     — table with question column; partially fillable
  * "docx_list"      — numbered/bulleted questions; export as fresh DOCX
  * "llm_freetext"   — caller (extractor) should run the LLM structurer
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from io import BytesIO
from typing import Any

log = logging.getLogger(__name__)


# ── Public data shapes ──────────────────────────────────────────────────────


@dataclass
class ParsedRequirement:
    text: str
    category: str | None = None
    external_ref: str | None = None
    # XLSX format-preservation breadcrumbs (None on non-XLSX).
    source_sheet: str | None = None
    source_row_index: int | None = None
    source_answer_column: str | None = None  # column letter ("C")
    notes_hint: str | None = None             # text already in a Notes cell


@dataclass
class ParsedRfp:
    requirements: list[ParsedRequirement]
    extraction_source: str
    # Map sheet → {question_column, answer_column, ref_column, category_column,
    #             notes_column, header_row_index}. Used by the exporter to
    # write answers back without re-discovering the structure.
    sheet_layout: dict[str, dict[str, Any]] = field(default_factory=dict)
    notes: str | None = None
    full_text_fallback: str | None = None  # used by extractor for LLM fallback


# ── Header-keyword heuristics ───────────────────────────────────────────────

_QUESTION_KEYWORDS = (
    "question", "requirement", "criteria", "criterion", "item",
    "description", "topic", "ask", "inquiry", "sub-requirement",
    "vendor question", "evaluation criteria",
)
_ANSWER_KEYWORDS = (
    "answer", "response", "vendor response", "supplier response",
    "vendor reply", "reply", "vendor answer", "vendor comments",
    "comply", "compliance", "y/n",
)
_REF_KEYWORDS = ("#", "id", "ref", "no.", "no", "number", "req id", "q#", "item #")
_CATEGORY_KEYWORDS = ("category", "section", "type", "domain", "module", "area")
_NOTES_KEYWORDS = ("notes", "comments", "evidence", "remarks", "explanation", "clarification")


def _matches(header: str, keywords: tuple[str, ...]) -> bool:
    h = (header or "").strip().lower()
    if not h:
        return False
    return any(k == h or k in h for k in keywords)


# ── Public entry point ─────────────────────────────────────────────────────


def parse_rfp(file_bytes: bytes, source_format: str) -> ParsedRfp:
    """Dispatch by file format. Never raises on extraction failure — returns
    a ParsedRfp with `full_text_fallback` so the LLM extractor can have a
    second go. Raises only on hard parse failure (file corrupt / unreadable).
    """
    fmt = source_format.lower()
    if fmt == "xlsx":
        return _parse_xlsx(file_bytes)
    if fmt == "docx":
        return _parse_docx(file_bytes)
    if fmt == "csv":
        return _parse_csv(file_bytes)
    # PDF / TXT / MD: no structure to lift. Defer to LLM extractor.
    return ParsedRfp(
        requirements=[],
        extraction_source="llm_freetext",
        full_text_fallback=_extract_plain_text(file_bytes, fmt),
    )


# ── XLSX ────────────────────────────────────────────────────────────────────


def _parse_xlsx(file_bytes: bytes) -> ParsedRfp:
    import openpyxl
    from openpyxl.utils import get_column_letter

    try:
        wb = openpyxl.load_workbook(BytesIO(file_bytes), data_only=True)
    except Exception as exc:
        raise RuntimeError(f"xlsx_open_failed: {exc}") from exc

    all_requirements: list[ParsedRequirement] = []
    sheet_layout: dict[str, dict[str, Any]] = {}
    full_text_chunks: list[str] = []

    try:
        for sheet in wb.worksheets:
            sheet_title = sheet.title or "Sheet"
            layout = _detect_xlsx_layout(sheet)
            if not layout:
                # Capture text for LLM fallback even if no question col found.
                full_text_chunks.append(_dump_sheet_text(sheet, sheet_title))
                continue

            sheet_layout[sheet_title] = layout
            q_idx = layout["question_column_index"]
            a_idx = layout["answer_column_index"]
            ref_idx = layout["ref_column_index"]
            cat_idx = layout["category_column_index"]
            notes_idx = layout["notes_column_index"]
            header_row = layout["header_row_index"]

            # Iterate rows after the header. Use max_row+1 cap.
            for row_idx, row in enumerate(
                sheet.iter_rows(min_row=header_row + 1, values_only=True),
                start=header_row + 1,
            ):
                cells = list(row)
                q_text = _cell_str(cells[q_idx]) if q_idx is not None and q_idx < len(cells) else ""
                if not q_text or len(q_text) < 5:
                    # Skip empty rows + obvious filler ("N/A", "TBD", etc.)
                    continue
                # Header rows can re-appear (section headers); detect and skip.
                if _looks_like_header_row(cells, layout):
                    continue
                req = ParsedRequirement(
                    text=q_text,
                    external_ref=_cell_str(cells[ref_idx]) if ref_idx is not None and ref_idx < len(cells) else None,
                    category=_cell_str(cells[cat_idx]) if cat_idx is not None and cat_idx < len(cells) else None,
                    notes_hint=_cell_str(cells[notes_idx]) if notes_idx is not None and notes_idx < len(cells) else None,
                    source_sheet=sheet_title,
                    source_row_index=row_idx,
                    source_answer_column=get_column_letter(a_idx + 1) if a_idx is not None else None,
                )
                all_requirements.append(req)
    finally:
        wb.close()

    if all_requirements:
        return ParsedRfp(
            requirements=all_requirements,
            extraction_source="xlsx_table",
            sheet_layout=sheet_layout,
        )

    # No structured rows lifted. Return text for LLM fallback.
    return ParsedRfp(
        requirements=[],
        extraction_source="llm_freetext",
        full_text_fallback="\n\n".join(full_text_chunks),
    )


def _detect_xlsx_layout(sheet) -> dict[str, Any] | None:
    """Scan the first 15 rows looking for the header row. A row is a header
    if at least 2 of its cells match known column-kind keywords (question +
    one other) OR if it has a question-like header.

    Returns the column layout, or None if no header could be detected.
    """
    candidate_rows = []
    for r_idx, row in enumerate(
        sheet.iter_rows(min_row=1, max_row=15, values_only=True), start=1
    ):
        headers = [(_cell_str(c) or "").strip() for c in row]
        if not any(headers):
            continue
        q_col = _find_column(headers, _QUESTION_KEYWORDS)
        if q_col is None:
            continue
        a_col = _find_column(headers, _ANSWER_KEYWORDS)
        ref_col = _find_column(headers, _REF_KEYWORDS)
        cat_col = _find_column(headers, _CATEGORY_KEYWORDS)
        notes_col = _find_column(headers, _NOTES_KEYWORDS)

        score = sum(c is not None for c in (q_col, a_col, ref_col, cat_col, notes_col))
        candidate_rows.append(
            (
                score,
                r_idx,
                {
                    "header_row_index": r_idx,
                    "header_labels": headers,
                    "question_column_index": q_col,
                    "answer_column_index": a_col,
                    "ref_column_index": ref_col,
                    "category_column_index": cat_col,
                    "notes_column_index": notes_col,
                },
            )
        )

    if not candidate_rows:
        return None
    # Highest score wins; tie → earlier row (more likely the real header).
    candidate_rows.sort(key=lambda x: (-x[0], x[1]))
    return candidate_rows[0][2]


def _find_column(headers: list[str], keywords: tuple[str, ...]) -> int | None:
    # Prefer exact match, then substring.
    norm = [h.strip().lower() for h in headers]
    for i, h in enumerate(norm):
        if h in keywords:
            return i
    for i, h in enumerate(norm):
        if h and any(k in h for k in keywords):
            return i
    return None


def _looks_like_header_row(cells: list[Any], layout: dict[str, Any]) -> bool:
    """A row that re-uses the header labels (section dividers in nested RFPs).
    If the question cell equals the header label literally, it's a divider."""
    q_idx = layout["question_column_index"]
    if q_idx is None or q_idx >= len(cells):
        return False
    val = (_cell_str(cells[q_idx]) or "").strip().lower()
    return val and val == (layout["header_labels"][q_idx] or "").strip().lower()


def _dump_sheet_text(sheet, sheet_title: str) -> str:
    parts = [f"## {sheet_title}"]
    for row in sheet.iter_rows(values_only=True, max_row=500):
        cells = [s for s in (_cell_str(c) for c in row) if s]
        if cells:
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _cell_str(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


# ── DOCX ────────────────────────────────────────────────────────────────────

_NUMBERED_RE = re.compile(r"^\s*(\d+(?:\.\d+)*)[.):]\s+(.+)$")
_LETTERED_RE = re.compile(r"^\s*([A-Za-z])[.):]\s+(.+)$")
_BULLET_RE = re.compile(r"^\s*[•\-\*]\s+(.+)$")


def _parse_docx(file_bytes: bytes) -> ParsedRfp:
    from docx import Document

    try:
        doc = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise RuntimeError(f"docx_open_failed: {exc}") from exc

    table_reqs = _extract_docx_tables(doc)
    if table_reqs:
        return ParsedRfp(
            requirements=table_reqs,
            extraction_source="docx_table",
        )

    list_reqs = _extract_docx_lists(doc)
    if list_reqs:
        return ParsedRfp(
            requirements=list_reqs,
            extraction_source="docx_list",
        )

    # No structure detected — fall back to LLM extraction on the prose.
    full_text = "\n\n".join(
        p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()
    )
    return ParsedRfp(
        requirements=[],
        extraction_source="llm_freetext",
        full_text_fallback=full_text,
    )


def _extract_docx_tables(doc) -> list[ParsedRequirement]:
    """Lift Q/A from DOCX tables. We pick the first table whose header row
    matches our keywords; rare RFPs have multiple tables but we treat them
    additively when each has a matching header."""
    out: list[ParsedRequirement] = []
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [c.text.strip() for c in table.rows[0].cells]
        q_col = _find_column(headers, _QUESTION_KEYWORDS)
        if q_col is None:
            continue
        ref_col = _find_column(headers, _REF_KEYWORDS)
        cat_col = _find_column(headers, _CATEGORY_KEYWORDS)
        notes_col = _find_column(headers, _NOTES_KEYWORDS)
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if q_col >= len(cells):
                continue
            q_text = cells[q_col]
            if not q_text or len(q_text) < 5:
                continue
            out.append(
                ParsedRequirement(
                    text=q_text,
                    external_ref=cells[ref_col] if ref_col is not None and ref_col < len(cells) else None,
                    category=cells[cat_col] if cat_col is not None and cat_col < len(cells) else None,
                    notes_hint=cells[notes_col] if notes_col is not None and notes_col < len(cells) else None,
                )
            )
    return out


def _extract_docx_lists(doc) -> list[ParsedRequirement]:
    """Numbered/bulleted question lists. Uses two signals: paragraph numbering
    style (when python-docx exposes it) and a regex over the visible text."""
    out: list[ParsedRequirement] = []
    current_category: str | None = None

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = (para.style.name or "").lower() if para.style else ""
        if style_name.startswith("heading") or style_name in ("title", "subtitle"):
            current_category = text[:120]
            continue

        m = _NUMBERED_RE.match(text)
        if m:
            ref, body = m.group(1), m.group(2).strip()
            if len(body) >= 5:
                out.append(
                    ParsedRequirement(
                        text=body,
                        external_ref=ref,
                        category=current_category,
                    )
                )
            continue
        m = _LETTERED_RE.match(text)
        if m:
            ref, body = m.group(1).upper(), m.group(2).strip()
            if len(body) >= 5 and body.endswith("?"):
                out.append(
                    ParsedRequirement(
                        text=body,
                        external_ref=ref,
                        category=current_category,
                    )
                )
            continue
        m = _BULLET_RE.match(text)
        if m:
            body = m.group(1).strip()
            if len(body) >= 5 and ("?" in body or len(body) > 30):
                out.append(
                    ParsedRequirement(text=body, category=current_category)
                )
            continue

    return out


# ── CSV ─────────────────────────────────────────────────────────────────────


def _parse_csv(file_bytes: bytes) -> ParsedRfp:
    import csv
    import io

    text = _decode(file_bytes)
    try:
        dialect = csv.Sniffer().sniff(text[:8192], delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    reader = csv.reader(io.StringIO(text), dialect=dialect)

    rows = list(reader)
    if not rows:
        return ParsedRfp(requirements=[], extraction_source="llm_freetext", full_text_fallback=text)

    headers = [c.strip() for c in rows[0]]
    q_col = _find_column(headers, _QUESTION_KEYWORDS)
    if q_col is None:
        return ParsedRfp(requirements=[], extraction_source="llm_freetext", full_text_fallback=text)

    ref_col = _find_column(headers, _REF_KEYWORDS)
    cat_col = _find_column(headers, _CATEGORY_KEYWORDS)
    notes_col = _find_column(headers, _NOTES_KEYWORDS)

    out: list[ParsedRequirement] = []
    for row in rows[1:]:
        cells = [c.strip() for c in row]
        if q_col >= len(cells):
            continue
        q_text = cells[q_col]
        if not q_text or len(q_text) < 5:
            continue
        out.append(
            ParsedRequirement(
                text=q_text,
                external_ref=cells[ref_col] if ref_col is not None and ref_col < len(cells) else None,
                category=cells[cat_col] if cat_col is not None and cat_col < len(cells) else None,
                notes_hint=cells[notes_col] if notes_col is not None and notes_col < len(cells) else None,
            )
        )
    # CSV → exporter cannot round-trip into the source file; treat like docx_table.
    return ParsedRfp(requirements=out, extraction_source="docx_table" if out else "llm_freetext")


def _decode(file_bytes: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def _extract_plain_text(file_bytes: bytes, fmt: str) -> str:
    """Lightweight text extraction for LLM fallback. Reuses ingestion.parser
    for PDF (PyMuPDF). Other formats: best-effort decode."""
    if fmt == "pdf":
        try:
            from app.services.ingestion.parser import parse_document
            segs = parse_document(file_bytes, "pdf")
            return "\n\n".join(s.content for s in segs if (s.content or "").strip())
        except Exception as exc:
            log.warning("rfp.parser.pdf_text_extract_failed: %s", exc)
            return ""
    if fmt in ("txt", "md"):
        return _decode(file_bytes)
    return ""
