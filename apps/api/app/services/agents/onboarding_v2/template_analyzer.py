"""AI-assisted DOCX template analysis — classification and in-place editing.

HR teams almost never upload DOCX templates with Jinja `{{ }}` placeholders
already in them: they have a Word document with blank spaces (underscores,
bracketed labels, empty table cells) they used to fill in by hand. Bridging
that gap is a three-stage job, and this module owns the middle one:

  1. **Detect** where values go — `blank_detector.find_all_candidates`.
     Deterministic, local, no LLM. Also `jinja_validator.classify_template`,
     which decides whether a document is a genuine hand-authored template.

  2. **Classify** each detected fill-point against the fixed vocabulary in
     `template_vars.py` — `classify_candidates` below. One batched LLM call per
     template (never per generated document), seeing only ~60 chars of context
     per fill-point and a closed list of variable names. It cannot invent
     variables and never reproduces or rewrites document text, so there is no
     verbatim-echo step that could drift.

  3. **Fill** the template — `pdf.slot_renderer`, from the `template_field_slots`
     rows HR confirmed. Nothing is ever written into the customer's document, so
     it never has to be valid template-language source.

Stage 3 used to have a step before it that rewrote the customer's `.docx` to
contain `{{ variable }}` tags so docxtpl could render it. That step is gone
(migration 097): a document only reaches docxtpl now if HR hand-authored it with
valid tags, which is what `validate_rendered` below still guards.

`extract_editable_blocks` / `apply_text_edits` are orthogonal to all of it: they
back the flat-text editor HR uses to reword a template in place, and are keyed
off the same canonical paragraph enumeration (`docx_positions`) as everything
else.
"""
from __future__ import annotations

import io
import json
import logging
import re
from dataclasses import dataclass
from typing import Any

from docx import Document  # python-docx
from docx.text.paragraph import Paragraph

from app.services.llm.client import LLMError, get_llm_client
from app.services.llm.types import Message
from app.services.pdf import (
    PdfRenderError,
    PdfRenderUnavailable,
    TemplateVariableError,
    fill_docx_template,
)

from .blank_detector import (
    BlankCandidate,
    build_candidate_block,
    find_all_candidates,
)
from .docx_positions import canonical_paragraphs
from .jinja_validator import classify_template
from .render_context import sample_render_context
from .template_vars import TEMPLATE_VARIABLES, get_variable_names

__all__ = [
    "BlankCandidate",
    "CandidateClassification",
    "EditableBlock",
    "TemplateAnalyzerError",
    "apply_text_edits",
    "classify_candidates",
    "classify_template",
    "extract_editable_blocks",
    "extract_text",
    "find_all_candidates",
    "validate_rendered",
]

log = logging.getLogger(__name__)


# ── Public types ───────────────────────────────────────────────────────────


@dataclass(frozen=True)
class CandidateClassification:
    """The LLM's decision for one candidate: which variable fills it, or None
    to skip it (no vocabulary variable fits)."""

    candidate_id: int
    variable: str | None
    confidence: str  # "high" | "medium" | "low"


class TemplateAnalyzerError(RuntimeError):
    """Raised when the analyzer can't produce usable output. The router
    converts this to a 422 with a human-readable detail so HR sees an
    actionable error rather than a 500."""


# ── Text extraction ────────────────────────────────────────────────────────


def extract_text(docx_bytes: bytes) -> str:
    """Extract plain text from a DOCX, preserving paragraph + table order.

    Used to (a) detect existing Jinja placeholders and (b) feed the LLM a
    flat view of the document. Header/footer text is included since LOI/AL
    templates often put the candidate name in a letterhead block.
    """
    doc = Document(io.BytesIO(docx_bytes))
    lines: list[str] = []

    def _add(text: str) -> None:
        if text and text.strip():
            lines.append(text)

    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                _add(p.text)
        if section.footer:
            for p in section.footer.paragraphs:
                _add(p.text)

    for block in _iter_body_blocks(doc):
        _add(block)

    return "\n".join(lines)


def _iter_body_blocks(doc: Any) -> list[str]:
    """Walk the body in document order so paragraphs and tables interleave
    correctly. python-docx exposes them as separate lists by default — we
    inspect the underlying XML order to get the linear sequence."""
    from docx.oxml.ns import qn

    out: list[str] = []
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            # Locate the matching Paragraph object so we get its full text
            # (including style/run merging) instead of re-doing it from XML.
            para = Paragraph(child, doc)
            out.append(para.text)
        elif child.tag == qn("w:tbl"):
            for row in child.iter(qn("w:tr")):
                cells_text: list[str] = []
                for cell in row.iter(qn("w:tc")):
                    cell_text = " ".join(
                        p.text for p in (Paragraph(p_elem, doc) for p_elem in cell.iter(qn("w:p")))
                    ).strip()
                    if cell_text:
                        cells_text.append(cell_text)
                if cells_text:
                    out.append(" | ".join(cells_text))
    return out


# ── LLM mapping proposal ───────────────────────────────────────────────────


_CLASSIFIER_SYSTEM_PROMPT = """You are a document-template analyzer for HR documents (Letter of Intent, Appointment Letter, NDA, or Induction).

You are given:
1. A fixed vocabulary of variables that can be substituted into the template.
2. A numbered list of BLANK CANDIDATES already located in the document. Each
   candidate is a real blank fill-spot (underscores, a bracketed label, etc.)
   shown with its surrounding text; the blank itself is wrapped in ⟦ ⟧ markers.

Your ONLY job: for EVERY candidate, decide which vocabulary variable belongs in
that blank — or that none does. You do NOT locate blanks (that is already done)
and you do NOT reproduce or rewrite any document text.

Rules:
  * Pick the `variable` from the vocabulary that should fill the blank, judging
    from the surrounding text — labels like "Name:", "Date:", "CTC:", a
    signature line, the governing-law clause, etc.
  * If no vocabulary variable fits a blank, set `variable` to null. Do NOT
    invent a variable outside the vocabulary.
  * Return EXACTLY one entry for every candidate_id you were given — never omit
    one and never add ids that weren't given.
  * `confidence` is how sure you are: "high" | "medium" | "low". Low is fine —
    HR reviews every mapping before it is applied.

Output format — STRICT JSON only, no prose, no markdown fences:

{
  "classifications": [
    { "candidate_id": <int>, "variable": "<vocabulary name or null>", "confidence": "high" | "medium" | "low" }
  ]
}

Return ONLY the JSON object. No explanation."""


def _build_vocabulary_block() -> str:
    """Render the variable vocabulary for inclusion in the LLM user prompt."""
    lines = ["Vocabulary (the only allowed `variable` values):"]
    for v in TEMPLATE_VARIABLES:
        lines.append(f"  - {v['name']}: {v['description']}")
    return "\n".join(lines)


async def classify_candidates(
    *,
    candidates: list[BlankCandidate],
    template_kind: str,
) -> list[CandidateClassification]:
    """Ask the LLM to assign a vocabulary variable (or null=skip) to each
    pre-located blank.

    This is a constrained classification task: the model only picks a variable
    name per `candidate_id` and never reproduces document text, so there is no
    verbatim-echo step that can drift. Raises TemplateAnalyzerError if the LLM
    is unreachable or returns unparseable JSON.

    Candidates the model omits or marks null simply get no entry here. They are
    NOT dropped downstream: `template_slots.build_slot_rows` still persists them
    with `variable=None` so HR sees the field and picks a value, because a blank
    the classifier didn't understand is exactly the one that must not go
    invisible.
    """
    if not candidates:
        return []

    client = get_llm_client()
    user_prompt = (
        f"Template kind: {template_kind}\n\n"
        f"{_build_vocabulary_block()}\n\n"
        f"{build_candidate_block(candidates)}\n\n"
        "Return the JSON classifications object now."
    )

    try:
        response = await client.complete(
            messages=[Message(role="user", content=user_prompt)],
            tools=(),
            temperature=0.0,
            replace_system_prompt=True,
            system_extra=_CLASSIFIER_SYSTEM_PROMPT,
        )
    except LLMError as exc:
        raise TemplateAnalyzerError(
            f"AI analyzer unavailable: {exc}. Try again or add {{{{ variable }}}} placeholders manually."
        ) from exc

    raw = (response.text or "").strip()
    if not raw:
        raise TemplateAnalyzerError("AI analyzer returned no output. Try again.")

    payload = _extract_json_object(raw)
    if payload is None:
        log.warning("template_analyzer.unparseable_response raw=%r", raw[:500])
        raise TemplateAnalyzerError(
            "AI analyzer returned unparseable output. Try again or add placeholders manually."
        )

    raw_items = payload.get("classifications")
    if not isinstance(raw_items, list):
        raise TemplateAnalyzerError("AI analyzer returned malformed classifications.")

    known_ids = {c.candidate_id for c in candidates}
    allowed = set(get_variable_names())
    result: list[CandidateClassification] = []
    seen: set[int] = set()
    for entry in raw_items:
        if not isinstance(entry, dict):
            continue
        try:
            cid = int(entry.get("candidate_id"))
        except (TypeError, ValueError):
            continue
        if cid not in known_ids or cid in seen:
            continue
        var_raw = entry.get("variable")
        variable = str(var_raw).strip() if var_raw not in (None, "") else None
        if variable is not None and variable not in allowed:
            # Model named a variable outside the vocabulary — treat as skip
            # rather than invent one.
            log.info(
                "template_analyzer.skipping_unknown_var candidate_id=%s variable=%s",
                cid, variable,
            )
            variable = None
        confidence = str(entry.get("confidence") or "medium").lower()
        if confidence not in ("high", "medium", "low"):
            confidence = "medium"
        seen.add(cid)
        result.append(
            CandidateClassification(candidate_id=cid, variable=variable, confidence=confidence)
        )

    return result


def _extract_json_object(text: str) -> dict[str, Any] | None:
    """LLMs occasionally wrap JSON in ``` fences despite being told not to.
    Tolerate both bare-JSON and fenced output. Returns None if nothing
    valid parses."""
    cleaned = text.strip()
    if cleaned.startswith("```"):
        # Strip ```json ... ``` or ``` ... ```
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        parsed = json.loads(cleaned)
    except json.JSONDecodeError:
        # Last-ditch: extract the first {...} block.
        match = re.search(r"\{.*\}", cleaned, re.DOTALL)
        if not match:
            return None
        try:
            parsed = json.loads(match.group(0))
        except json.JSONDecodeError:
            return None
    return parsed if isinstance(parsed, dict) else None


# ── Flat-text editing (paragraph-in-place) ─────────────────────────────────
#
# HR reviews a template as flat, editable text — one field per paragraph /
# table cell / header / footer line. When they save, we match each edited
# line back to its original DOCX paragraph *by position* and rewrite the run
# text in place. This preserves fonts, tables, letterheads, and existing
# `{{ placeholders }}`; HR can reword freely but can't add/remove/reorder
# paragraphs inline (that still goes through Download → edit → Replace .docx).


@dataclass(frozen=True)
class EditableBlock:
    """One editable line of a template.

    `index` is the position in the canonical paragraph enumeration — it is the
    stable key the write-back step uses to locate the paragraph again, so the
    SAME enumeration must drive both extraction and apply. `kind` is a display
    hint ("body" | "table" | "header" | "footer").
    """

    index: int
    text: str
    kind: str


def extract_editable_blocks(docx_bytes: bytes) -> list[EditableBlock]:
    """Return the non-empty paragraphs of a DOCX as editable blocks.

    Blank/spacing paragraphs are skipped from the editable list but keep their
    slot in the canonical enumeration, so every block's `index` still points
    at the right paragraph on write-back.
    """
    doc = Document(io.BytesIO(docx_bytes))
    blocks: list[EditableBlock] = []
    for idx, (para, kind) in enumerate(canonical_paragraphs(doc)):
        text = para.text
        if text and text.strip():
            blocks.append(EditableBlock(index=idx, text=text, kind=kind))
    return blocks


def _set_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    """Overwrite a paragraph's text, preserving its first run's formatting.

    Word splits styled text across runs; we write the whole new string into
    the first run and blank the rest. Sub-run formatting on the replaced span
    collapses to the first run's style — acceptable for paragraph-level editing.
    """
    runs = paragraph.runs
    if not runs:
        # No runs to inherit style from (shouldn't happen for a non-empty
        # block, but guard anyway) — add a plain run.
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def apply_text_edits(*, docx_bytes: bytes, edits: dict[int, str]) -> tuple[bytes, int]:
    """Rewrite paragraphs at the given canonical indices with new text.

    `edits` maps a block `index` (from `extract_editable_blocks`) to its new
    text. Indices out of range are ignored; paragraphs whose text is unchanged
    are left untouched. Returns the new DOCX bytes and the number of
    paragraphs actually changed.
    """
    doc = Document(io.BytesIO(docx_bytes))
    paras = canonical_paragraphs(doc)
    changed = 0
    for idx, new_text in edits.items():
        if idx < 0 or idx >= len(paras):
            continue
        para, _kind = paras[idx]
        if para.text == new_text:
            continue
        _set_paragraph_text(para, new_text)
        changed += 1

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), changed


# ── Validation ─────────────────────────────────────────────────────────────


async def validate_rendered(docx_bytes: bytes, *, template_kind: str | None = None) -> None:
    """Dry-render an EDITED docxtpl template with sample context, to catch a
    broken edit before we persist it.

    Only meaningful for `fill_strategy='jinja'` (and legacy NULL) templates —
    the ones docxtpl actually parses. Callers must skip it for slots templates,
    where running the document through Jinja would reject it for containing text
    that merely LOOKS like a placeholder. Raises TemplateAnalyzerError on render
    failure with the underlying message.
    """
    try:
        await fill_docx_template(
            template_bytes=docx_bytes,
            context=sample_render_context(),
            strict=True,
            template_kind=template_kind,
        )
    except TemplateVariableError as exc:
        raise TemplateAnalyzerError(
            f"This template references a placeholder we don't supply: "
            f"'{exc.variable_name}'. Edit the wording to remove it, or rename it "
            "to one of the available variables."
        ) from exc
    except (PdfRenderError, PdfRenderUnavailable) as exc:
        raise TemplateAnalyzerError(
            f"Couldn't render the edited template: {exc}"
        ) from exc
