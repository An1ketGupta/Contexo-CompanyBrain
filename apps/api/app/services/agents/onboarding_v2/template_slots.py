"""Persistence for template fill-points (`template_field_slots`).

A *slot* is one place in a template where a value goes, addressed by position:
`(paragraph_index, start_offset, end_offset)` in `canonical_paragraphs()` order,
plus the `action` describing how to write there and the `variable` that fills it.

This table is what replaced rewriting the customer's `.docx` to contain
`{{ tags }}`. Because the mapping lives here instead of in the document, the
document never has to be valid template-language source — and a malformed
hand-typed placeholder can no longer take down a render.

Two rules the whole review workflow rests on:

  * **Re-analyzing never undoes a human decision.** `upsert_proposed_slots`
    writes only to rows still `status='proposed'`. A slot HR confirmed or
    rejected keeps its decision even if the classifier changes its mind on a
    later pass.

  * **Only `confirmed` slots render.** A `proposed` slot is a suggestion nobody
    has approved yet; writing it into a legally-binding document would defeat
    the point of asking. `list_slots(confirmed_only=True)` is what the agent
    calls.

All writes go through the service-role client, matching the RLS posture in
migration 096 (select-only policy for end users).
"""
from __future__ import annotations

import asyncio
import logging
from datetime import UTC, datetime
from typing import Any

from app.database import get_service_client

from .blank_detector import BlankCandidate
from .docx_positions import canonical_paragraphs, docx_hash, paragraph_hash, paragraph_run_text

log = logging.getLogger(__name__)

_SLOT_COLUMNS = (
    "id, org_id, document_id, document_version_id, paragraph_index, "
    "paragraph_kind, paragraph_hash, source_docx_sha256, start_offset, "
    "end_offset, action, blank_text, context_before, context_after, "
    "variable, confidence, status, source, confirmed_by, confirmed_at, "
    "created_at, updated_at"
)


def build_slot_rows(
    *,
    org_id: str,
    document_id: str,
    document_version_id: str | None,
    docx_bytes: bytes,
    candidates: list[BlankCandidate],
    variables_by_candidate: dict[int, str | None],
    confidence_by_candidate: dict[int, str],
) -> list[dict[str, Any]]:
    """Turn detected candidates + their classifications into insertable rows.

    Hashes each candidate's anchor paragraph here, at detection time, from the
    same bytes the offsets were computed against — that pairing is what makes
    the render-time drift check meaningful.
    """
    import io

    from docx import Document  # local import: keeps module import cheap

    doc = Document(io.BytesIO(docx_bytes))
    canonical = canonical_paragraphs(doc)
    file_hash = docx_hash(docx_bytes)

    rows: list[dict[str, Any]] = []
    for cand in candidates:
        if cand.paragraph_index >= len(canonical):
            continue
        paragraph, _kind = canonical[cand.paragraph_index]
        rows.append(
            {
                "org_id": org_id,
                "document_id": document_id,
                "document_version_id": document_version_id,
                "paragraph_index": cand.paragraph_index,
                "paragraph_kind": cand.paragraph_kind,
                "paragraph_hash": paragraph_hash(paragraph_run_text(paragraph)),
                "source_docx_sha256": file_hash,
                "start_offset": cand.start_offset,
                "end_offset": cand.end_offset,
                "action": cand.action,
                "blank_text": cand.matched_text,
                "context_before": cand.context_before,
                "context_after": cand.context_after,
                "variable": variables_by_candidate.get(cand.candidate_id),
                "confidence": confidence_by_candidate.get(cand.candidate_id, "medium"),
                "status": "proposed",
                "source": "ai",
            }
        )
    return rows


async def upsert_proposed_slots(
    *,
    org_id: str,
    document_id: str,
    rows: list[dict[str, Any]],
) -> tuple[int, int]:
    """Persist freshly-detected proposals, preserving every HR decision.

    Returns `(inserted, refreshed)`. An anchor HR has already confirmed or
    rejected is left completely untouched — which is what makes "re-run
    Find fields" a safe button rather than a destructive one.

    Also prunes stale `proposed` rows for this document that the current
    detection pass no longer finds (e.g. HR deleted that blank from the
    template), so the review list never shows fields that aren't there any
    more. Confirmed/rejected rows are never pruned; drift detection is what
    speaks to those.
    """
    svc = get_service_client()

    existing = await list_slots(document_id=document_id)
    decided = {
        (r["paragraph_index"], r["start_offset"], r["end_offset"])
        for r in existing
        if r["status"] in ("confirmed", "rejected")
    }
    incoming_anchors = {
        (r["paragraph_index"], r["start_offset"], r["end_offset"]) for r in rows
    }

    fresh = [
        r
        for r in rows
        if (r["paragraph_index"], r["start_offset"], r["end_offset"]) not in decided
    ]

    stale_ids = [
        r["id"]
        for r in existing
        if r["status"] == "proposed"
        and (r["paragraph_index"], r["start_offset"], r["end_offset"]) not in incoming_anchors
    ]

    def _write() -> None:
        if stale_ids:
            svc.table("template_field_slots").delete().in_("id", stale_ids).execute()
        if fresh:
            # on_conflict targets uniq_template_field_slots_anchor — a
            # re-detected proposal refreshes its hash/context/variable in place.
            (
                svc.table("template_field_slots")
                .upsert(
                    fresh,
                    on_conflict="document_id,paragraph_index,start_offset,end_offset",
                )
                .execute()
            )

    await asyncio.to_thread(_write)

    if stale_ids:
        log.info(
            "template_slots.pruned_stale org=%s doc=%s count=%s",
            org_id, document_id, len(stale_ids),
        )
    return len(fresh), len(stale_ids)


async def list_slots(
    *,
    document_id: str,
    confirmed_only: bool = False,
) -> list[dict[str, Any]]:
    """All slots for a template, in document order.

    `confirmed_only=True` is the render-time view — the agent must never see a
    proposal nobody approved.
    """
    svc = get_service_client()

    def _read() -> list[dict[str, Any]]:
        q = (
            svc.table("template_field_slots")
            .select(_SLOT_COLUMNS)
            .eq("document_id", document_id)
        )
        if confirmed_only:
            q = q.eq("status", "confirmed")
        res = (
            q.order("paragraph_index", desc=False)
            .order("start_offset", desc=False)
            .execute()
        )
        return res.data or []

    return await asyncio.to_thread(_read)


async def count_pending(*, document_id: str) -> int:
    """How many proposals still await HR review. Gates promoting a template to
    active — an unreviewed field is a field that silently won't be filled."""
    svc = get_service_client()

    def _read() -> int:
        res = (
            svc.table("template_field_slots")
            .select("id", count="exact")
            .eq("document_id", document_id)
            .eq("status", "proposed")
            .execute()
        )
        return res.count or 0

    return await asyncio.to_thread(_read)


async def set_slot_status(
    *,
    slot_id: str,
    org_id: str,
    status: str,
    variable: str | None = None,
    user_id: str | None = None,
) -> dict[str, Any] | None:
    """Record HR's decision on one slot.

    `org_id` is part of the WHERE clause, not just a lookup: these writes use
    the service-role client, which bypasses RLS, so tenant scoping has to be
    explicit here or a slot id from another org would be editable.
    """
    svc = get_service_client()

    patch: dict[str, Any] = {"status": status}
    if variable is not None:
        patch["variable"] = variable
    if status == "confirmed":
        patch["confirmed_by"] = user_id
        patch["confirmed_at"] = datetime.now(UTC).isoformat()
    else:
        patch["confirmed_by"] = None
        patch["confirmed_at"] = None

    def _write() -> dict[str, Any] | None:
        res = (
            svc.table("template_field_slots")
            .update(patch)
            .eq("id", slot_id)
            .eq("org_id", org_id)
            .execute()
        )
        return (res.data or [None])[0]

    return await asyncio.to_thread(_write)


async def insert_manual_slot(
    *,
    org_id: str,
    document_id: str,
    document_version_id: str | None,
    docx_bytes: bytes,
    paragraph_index: int,
    start_offset: int,
    end_offset: int,
    action: str,
    variable: str,
    user_id: str | None = None,
) -> dict[str, Any]:
    """Add a fill-point HR located by hand.

    This is the recall safety net. Automated detection will never be perfect on
    arbitrary customer documents, so HR must be able to say "there is also a
    field here" — otherwise every undetectable blank becomes a silent gap in a
    generated contract, which is the failure mode this whole pipeline exists to
    remove. Inserted `confirmed`, since HR pointing at it *is* the confirmation.

    Raises ValueError if the anchor doesn't resolve in the current bytes, so a
    stale UI can't create a slot pointing at nothing.
    """
    import io

    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    canonical = canonical_paragraphs(doc)
    if paragraph_index < 0 or paragraph_index >= len(canonical):
        raise ValueError(f"paragraph_index {paragraph_index} is out of range")

    paragraph, kind = canonical[paragraph_index]
    text = paragraph_run_text(paragraph)
    if start_offset < 0 or end_offset > len(text) or start_offset > end_offset:
        raise ValueError(
            f"offsets [{start_offset}:{end_offset}] don't fit paragraph "
            f"{paragraph_index} (length {len(text)})"
        )

    row = {
        "org_id": org_id,
        "document_id": document_id,
        "document_version_id": document_version_id,
        "paragraph_index": paragraph_index,
        "paragraph_kind": kind,
        "paragraph_hash": paragraph_hash(text),
        "source_docx_sha256": docx_hash(docx_bytes),
        "start_offset": start_offset,
        "end_offset": end_offset,
        "action": action,
        "blank_text": text[start_offset:end_offset],
        "context_before": text[max(0, start_offset - 60):start_offset],
        "context_after": text[end_offset:end_offset + 60],
        "variable": variable,
        "confidence": "high",
        "status": "confirmed",
        "source": "manual",
        "confirmed_by": user_id,
        "confirmed_at": datetime.now(UTC).isoformat(),
    }

    svc = get_service_client()

    def _write() -> dict[str, Any]:
        res = (
            svc.table("template_field_slots")
            .upsert(
                row,
                on_conflict="document_id,paragraph_index,start_offset,end_offset",
            )
            .execute()
        )
        return (res.data or [row])[0]

    return await asyncio.to_thread(_write)


async def delete_slots_for_document(*, document_id: str) -> None:
    """Drop every slot for a template.

    Called when the underlying file is wholesale replaced (HR re-uploads an
    edited `.docx`): the old anchors describe a document that no longer exists,
    and keeping them would mean drift errors instead of a clean re-analyze.
    """
    svc = get_service_client()

    def _write() -> None:
        svc.table("template_field_slots").delete().eq("document_id", document_id).execute()

    await asyncio.to_thread(_write)


def check_drift(
    *,
    docx_bytes: bytes,
    slots: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Split slots into `(ok, drifted)` by re-hashing their anchor paragraphs.

    Deliberately per-paragraph rather than whole-file: HR reworded one clause in
    a six-page appointment letter should not invalidate the twelve fields
    elsewhere in it. Only slots whose own anchor paragraph changed are drifted.

    A slot whose `paragraph_index` no longer exists at all (the paragraph was
    deleted) counts as drifted, not as an error — same remedy either way.
    """
    import io

    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    canonical = canonical_paragraphs(doc)

    # Hash each touched paragraph once, not once per slot in it.
    hash_cache: dict[int, str | None] = {}

    ok: list[dict[str, Any]] = []
    drifted: list[dict[str, Any]] = []
    for slot in slots:
        idx = slot["paragraph_index"]
        if idx not in hash_cache:
            if 0 <= idx < len(canonical):
                hash_cache[idx] = paragraph_hash(paragraph_run_text(canonical[idx][0]))
            else:
                hash_cache[idx] = None
        current = hash_cache[idx]
        if current is not None and current == slot["paragraph_hash"]:
            ok.append(slot)
        else:
            drifted.append(slot)
    return ok, drifted
