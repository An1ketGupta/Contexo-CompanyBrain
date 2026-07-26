"""Locate the fill-points in an HR-authored DOCX.

Detection runs entirely locally and deterministically — no LLM — and is
deliberately biased toward RECALL. Over-flagging is cheap: a false positive
becomes a candidate the classifier maps to `None`, or one HR rejects in a
single click. Under-flagging is what actually ships a contract with a blank
signature line, so a missed fill-point is the expensive failure.

Three heuristics, because HR documents mark "type here" in three different
ways and the first one alone was demonstrably not enough:

  1. `find_blank_candidates` — a visible marker span: `________`,
     `[CANDIDATE NAME]`, `<NAME>`, a dotted leader, `{name}`. This was the
     original (and only) detector.

  2. `find_label_colon_candidates` — a label with *nothing after it*:
     a paragraph that is exactly `Employee Signature:` and then ends. There is
     no marker character to match here, which is why the regex-only detector
     was blind to it: no candidate, no error, and a document that shipped with
     the field silently unfilled.

  3. `find_empty_cell_candidates` — an empty table cell beside a labelled one,
     the two-column HR form layout. Also invisible to a span regex, for the
     same reason: the fill-point is an absence of text, not a shape of text.

The output of all three is a `BlankCandidate` addressed by
`(paragraph_index, start_offset, end_offset)` in `canonical_paragraphs()`
order, so the classifier, the persisted slot row, and the renderer all point at
the same place. Only stage 2 (classification) costs LLM tokens, and it only
ever sees this shortlist — token spend scales with the number of fill-points,
not the length of the document.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Any

from docx import Document  # python-docx
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from .docx_positions import canonical_paragraphs, paragraph_run_text

# ── Actions ────────────────────────────────────────────────────────────────
# What the renderer must DO at a fill-point. Mirrors the CHECK constraint on
# template_field_slots.action (migration 096).

ACTION_REPLACE_SPAN = "replace_span"
ACTION_INSERT_AFTER_LABEL = "insert_after_label"
ACTION_INSERT_EMPTY_CELL = "insert_empty_cell"

# ── Patterns ───────────────────────────────────────────────────────────────

# Any existing `{{ ... }}` span. Used to reject blank candidates that fall
# inside a placeholder in a hand-authored template.
JINJA_SPAN_RE = re.compile(r"\{\{.*?\}\}", re.DOTALL)

# A blank fill-spot marked by visible characters, matched as a SPAN so we get
# exact offsets. Ordered alternation: underscore runs, bracketed labels,
# angle-bracket labels, dotted leaders / ellipsis, and single-curly `{label}`
# (the lookarounds exclude `{{ }}` Jinja, handled separately).
_BLANK_SPAN_RE = re.compile(
    r"_{3,}"                            # ________
    r"|\[[^\[\]{}\n]{1,60}\]"           # [CANDIDATE NAME]
    r"|<[^<>{}\n]{1,60}>"               # <NAME>
    r"|\.{3,}"                          # dotted leader ......
    r"|…"                               # unicode ellipsis
    r"|(?<!\{)\{[^{}\n]{1,60}\}(?!\})"  # {name} but never {{ jinja }}
)

# A label with nothing after it. Anchored to the whole (stripped) paragraph and
# capped at 60 chars so it fires on `Employee Signature:` but not on a prose
# sentence that happens to end in a colon. The leading-capital requirement is
# doing real work: it rejects mid-sentence lead-ins like `... as follows:`.
_LABEL_COLON_RE = re.compile(r"^[A-Z][A-Za-z0-9 /()&.'\-]{0,58}[:：]$")

# Colon-terminated prose lead-ins, not field labels. These pass the shape test
# above (short, capitalised, colon-terminated) but a value must never be
# injected after them. Matched as a SUFFIX, not an exact equality: the real
# documents say "The terms are as follows:", not a bare "As follows:".
_LEAD_IN_SUFFIXES = (
    "as follows",
    "the following",
    "following",
    "as under",
    "as below",
    "below",
    "for example",
    "namely",
    "provided that",
    "in witness whereof",
    "witnesseth",
    "whereas",
)

# Words that only appear in a sentence, never in a form field label. A single
# hit is enough to reject — "The terms are as follows:" is caught by `are` even
# if the phrasing isn't in the suffix list above. Deliberately excludes common
# label filler like "of"/"the" ("Name of the Employee:" is a real label).
_PROSE_WORDS = frozenset({
    "are", "is", "was", "were", "be", "been", "being",
    "will", "shall", "may", "must", "should", "would", "can",
    "has", "have", "had", "do", "does", "did",
    "include", "includes", "including", "means", "agrees", "hereby",
    "you", "your", "we", "our", "they", "their", "this", "these", "those",
})

# Cell labels that mark the neighbouring cell as a fill-point even without a
# trailing colon — the common headings in a two-column HR form table.
_CELL_LABEL_WORDS = frozenset({
    "name", "full name", "candidate name", "employee name",
    "date", "start date", "joining date", "date of joining",
    "signature", "signed", "place", "witness",
    "designation", "role", "title", "job title", "position",
    "department", "location", "work location",
    "ctc", "salary", "compensation",
    "email", "email id", "phone", "contact", "mobile",
    "manager", "reporting manager", "employee id", "id",
    "probation", "jurisdiction", "company",
})

# How much surrounding text to capture on each side of a fill-point — enough
# for the classifier (and HR) to read the label without bloating the prompt.
_CONTEXT_CHARS = 60


@dataclass(frozen=True)
class BlankCandidate:
    """A fill-point located deterministically by exact position.

    Detection never asks the LLM to reproduce document text — these offsets are
    the single source of truth for both the classification prompt and the later
    substitution. So nothing can drift, and two identical-looking fill-points
    (e.g. two `________` signature lines) are always distinguished by
    `paragraph_index` + `start_offset`, never by their (identical) text.

    For the two `insert_*` actions there is no span to replace:
    `start_offset == end_offset` marks an insertion point and `matched_text` is
    empty.
    """

    candidate_id: int
    paragraph_index: int  # index into canonical_paragraphs()
    start_offset: int     # char offset into the paragraph's run text
    end_offset: int
    matched_text: str     # exact substring — regex-extracted, never LLM-typed
    context_before: str
    context_after: str
    paragraph_kind: str   # "body" | "table" | "header" | "footer"
    action: str = ACTION_REPLACE_SPAN


# ── Marker-span detection (heuristic 1) ────────────────────────────────────


def _span_candidates(
    canonical: list[tuple[Paragraph, str]],
    *,
    skip_indices: set[int] | None = None,
) -> list[BlankCandidate]:
    """Locate every visibly-marked blank span, by exact character offset.

    Candidates overlapping an existing `{{ ... }}` span are dropped so we never
    re-template an already-templated field.
    """
    skip = skip_indices or set()
    out: list[BlankCandidate] = []
    for idx, (paragraph, kind) in enumerate(canonical):
        if idx in skip:
            continue
        text = paragraph_run_text(paragraph)
        if not text:
            continue
        jinja_spans = [(m.start(), m.end()) for m in JINJA_SPAN_RE.finditer(text)]
        for m in _BLANK_SPAN_RE.finditer(text):
            start, end = m.start(), m.end()
            if any(s < end and start < e for s, e in jinja_spans):
                continue  # inside/overlapping an existing {{ }} — not a blank
            out.append(
                BlankCandidate(
                    candidate_id=-1,  # assigned by find_all_candidates
                    paragraph_index=idx,
                    start_offset=start,
                    end_offset=end,
                    matched_text=m.group(0),
                    context_before=text[max(0, start - _CONTEXT_CHARS):start],
                    context_after=text[end:end + _CONTEXT_CHARS],
                    paragraph_kind=kind,
                    action=ACTION_REPLACE_SPAN,
                )
            )
    return out


def find_blank_candidates(docx_bytes: bytes) -> list[BlankCandidate]:
    """Marker-span fill-points only. Kept as a standalone entry point for
    tests and for callers that explicitly want just the original behaviour;
    production code should call `find_all_candidates`."""
    doc = Document(io.BytesIO(docx_bytes))
    return _renumber(_span_candidates(canonical_paragraphs(doc)))


# ── Label-with-nothing-after detection (heuristic 2) ───────────────────────


def _is_field_label(text: str) -> bool:
    """True if `text` reads as a form field label rather than prose.

    Three guards, all necessary. The shape test alone accepts
    `The terms are as follows:` — a clause that ends a sentence and must never
    have a candidate's name appended to it. The suffix list catches the common
    legal-document lead-ins, and the prose-word check catches the rest by
    grammar: a field label contains no verb.
    """
    stripped = text.strip()
    if not _LABEL_COLON_RE.match(stripped):
        return False
    bare = stripped[:-1].strip().lower()
    if not bare:
        return False
    if any(bare.endswith(suffix) for suffix in _LEAD_IN_SUFFIXES):
        return False
    words = bare.replace("/", " ").split()
    if any(w.strip(".,()") in _PROSE_WORDS for w in words):
        return False
    # A label is a handful of words, not a clause.
    return len(words) <= 6


def _label_colon_candidates(
    canonical: list[tuple[Paragraph, str]],
    *,
    skip_indices: set[int] | None = None,
) -> list[BlankCandidate]:
    """Locate labels that end in a colon with nothing after them.

    The insertion point is the END of the paragraph's run text, so the
    renderer appends the value in place, keeping the label's own formatting.
    """
    skip = skip_indices or set()
    out: list[BlankCandidate] = []
    for idx, (paragraph, kind) in enumerate(canonical):
        if idx in skip:
            continue
        text = paragraph_run_text(paragraph)
        if not text or not _is_field_label(text):
            continue
        if JINJA_SPAN_RE.search(text) or _BLANK_SPAN_RE.search(text):
            continue  # already templated, or already a marker-span candidate
        offset = len(text)
        out.append(
            BlankCandidate(
                candidate_id=-1,
                paragraph_index=idx,
                start_offset=offset,
                end_offset=offset,
                matched_text="",
                context_before=text[max(0, offset - _CONTEXT_CHARS):offset],
                context_after="",
                paragraph_kind=kind,
                action=ACTION_INSERT_AFTER_LABEL,
            )
        )
    return out


def find_label_colon_candidates(docx_bytes: bytes) -> list[BlankCandidate]:
    """Label-with-nothing-after fill-points only. Standalone for tests."""
    doc = Document(io.BytesIO(docx_bytes))
    return _renumber(_label_colon_candidates(canonical_paragraphs(doc)))


# ── Empty-table-cell detection (heuristic 3) ──────────────────────────────


def _iter_tables(container: Any) -> list[Table]:
    """All tables in a document, including tables nested inside cells."""
    found: list[Table] = []
    for table in getattr(container, "tables", []) or []:
        found.append(table)
        for row in table.rows:
            for cell in row.cells:
                found.extend(_iter_tables(cell))
    return found


def _cell_text(cell: _Cell) -> str:
    return " ".join(paragraph_run_text(p) for p in cell.paragraphs).strip()


def _is_cell_label(text: str) -> bool:
    bare = text.strip().rstrip(":：").strip().lower()
    if not bare or len(bare) > 60:
        return False
    return text.strip().endswith((":", "：")) or bare in _CELL_LABEL_WORDS


def _empty_cell_candidates(
    doc: Any,
    canonical: list[tuple[Paragraph, str]],
) -> tuple[list[BlankCandidate], set[int]]:
    """Locate empty cells sitting beside a labelled cell in the same row.

    Returns the candidates plus the canonical indices of the *label* paragraphs
    they consumed. `find_all_candidates` uses that second value to suppress the
    label-colon heuristic on those paragraphs — otherwise a `Name:` cell with
    an empty neighbour would yield two competing fill-points for one field.
    """
    index_by_element = {
        paragraph._element: idx for idx, (paragraph, _kind) in enumerate(canonical)
    }

    out: list[BlankCandidate] = []
    consumed_labels: set[int] = set()

    for table in _iter_tables(doc):
        for row in table.rows:
            try:
                cells = list(row.cells)
            except (IndexError, ValueError):
                # Malformed/merged row geometry — skip rather than guess.
                continue
            for left, right in zip(cells, cells[1:]):
                if left._tc is right._tc:
                    continue  # horizontally merged: same cell twice
                label_text = _cell_text(left)
                if _cell_text(right) or not _is_cell_label(label_text):
                    continue
                target = next(iter(right.paragraphs), None)
                if target is None:
                    continue
                idx = index_by_element.get(target._element)
                if idx is None:
                    continue
                out.append(
                    BlankCandidate(
                        candidate_id=-1,
                        paragraph_index=idx,
                        start_offset=0,
                        end_offset=0,
                        matched_text="",
                        context_before=label_text[-_CONTEXT_CHARS:],
                        context_after="",
                        paragraph_kind="table",
                        action=ACTION_INSERT_EMPTY_CELL,
                    )
                )
                for p in left.paragraphs:
                    label_idx = index_by_element.get(p._element)
                    if label_idx is not None:
                        consumed_labels.add(label_idx)

    return out, consumed_labels


def find_empty_cell_candidates(docx_bytes: bytes) -> list[BlankCandidate]:
    """Empty-cell-beside-a-label fill-points only. Standalone for tests."""
    doc = Document(io.BytesIO(docx_bytes))
    candidates, _consumed = _empty_cell_candidates(doc, canonical_paragraphs(doc))
    return _renumber(candidates)


# ── Combined entry point ───────────────────────────────────────────────────


def _renumber(candidates: list[BlankCandidate]) -> list[BlankCandidate]:
    """Assign sequential candidate_ids in document order.

    The classifier keys its response on `candidate_id`, so the ids must be
    dense and stable within one analyze call.
    """
    ordered = sorted(candidates, key=lambda c: (c.paragraph_index, c.start_offset))
    return [
        BlankCandidate(
            candidate_id=i,
            paragraph_index=c.paragraph_index,
            start_offset=c.start_offset,
            end_offset=c.end_offset,
            matched_text=c.matched_text,
            context_before=c.context_before,
            context_after=c.context_after,
            paragraph_kind=c.paragraph_kind,
            action=c.action,
        )
        for i, c in enumerate(ordered)
    ]


def find_all_candidates(docx_bytes: bytes) -> list[BlankCandidate]:
    """Run every heuristic over one parse of the document.

    Empty-cell detection runs first and claims its label paragraphs, so a
    `Name: | <empty>` table row produces exactly one fill-point (in the empty
    cell) rather than two competing ones.
    """
    doc = Document(io.BytesIO(docx_bytes))
    canonical = canonical_paragraphs(doc)

    cell_candidates, consumed_labels = _empty_cell_candidates(doc, canonical)
    occupied = {c.paragraph_index for c in cell_candidates}

    candidates = cell_candidates
    candidates += _span_candidates(canonical, skip_indices=occupied)
    candidates += _label_colon_candidates(
        canonical, skip_indices=occupied | consumed_labels
    )
    return _renumber(candidates)


# ── Classification prompt rendering ────────────────────────────────────────

_ACTION_HINTS = {
    ACTION_REPLACE_SPAN: "replaces-blank",
    ACTION_INSERT_AFTER_LABEL: "fills-after-label",
    ACTION_INSERT_EMPTY_CELL: "fills-empty-cell",
}


def build_candidate_block(candidates: list[BlankCandidate]) -> str:
    """Render the located candidates for the classification prompt — one line
    each, with the fill-point wrapped in ⟦ ⟧ inside its surrounding text.

    The action is included so the model can read an empty ⟦⟧ correctly: for the
    insert actions there is no blank text to look at, only the label before it.
    """
    lines = ["Blank candidates (classify EVERY one, by candidate_id):"]
    for c in candidates:
        snippet = f"{c.context_before}⟦{c.matched_text}⟧{c.context_after}"
        snippet = " ".join(snippet.split())  # collapse newlines/runs of spaces
        hint = _ACTION_HINTS.get(c.action, c.action)
        lines.append(f'  #{c.candidate_id} [{c.paragraph_kind}/{hint}]: "{snippet}"')
    return "\n".join(lines)


# ── Post-render safety net ─────────────────────────────────────────────────


def scan_for_unfilled_signals(text: str) -> list[str]:
    """Re-run the detection signals over ALREADY-RENDERED text.

    Detection will never have perfect recall on arbitrary customer documents,
    so the last line of defence is checking the output: if a rendered LOI still
    contains `________` or a bare `Signature:`, some fill-point was missed and
    HR should see that before the document reaches a candidate. Returns
    human-readable warnings (empty list = clean).

    Non-fatal by contract — the caller surfaces these as a warning, never as a
    render failure. A document with one missed field is still more useful to HR
    than no document at all.
    """
    warnings: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        match = _BLANK_SPAN_RE.search(line)
        if match:
            warnings.append(
                f"Still looks unfilled — '{match.group(0)}' in: {line[:80]}"
            )
            continue
        if _is_field_label(line):
            warnings.append(f"Label with no value after it: {line[:80]}")
    return warnings
