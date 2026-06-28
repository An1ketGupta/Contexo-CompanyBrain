"""Ship-with-product fallback templates for Onboarding v2.

Only NDA has a shipped default — LOI and Appointment Letter are too
jurisdiction-specific to ship a sensible default; we keep blocking the run
for those until HR uploads their own.

The default NDA is generated programmatically with python-docx so we don't
need to commit a binary .docx file (which would be brittle to review and
hard to diff). Customers can still customise by uploading their own NDA
and tagging it — the upload always wins over the default.

The placeholders below match the render context built by
OnboardingV2Agent._build_render_context — keep them in lockstep with
apps/api/tools/ONBOARDING_TEMPLATE_VARS.md.
"""
from __future__ import annotations

import asyncio
import io
import logging

log = logging.getLogger(__name__)


# Default NDA body — keep generic and balanced. We deliberately avoid
# overreaching IP-assignment clauses; this is a confidentiality agreement,
# not an IP transfer. Companies that need IP assignment should upload
# their own AL or addendum.
_NDA_TITLE = "Non-Disclosure Agreement"

_NDA_PARAGRAPHS: list[str] = [
    (
        "This Non-Disclosure Agreement (“Agreement”) is entered into on "
        "{{ today_date }} between {{ company_name }} (“Company”) and "
        "{{ candidate_name }} (“Receiving Party”), residing at the "
        "address on file with the Company."
    ),
    (
        "1. Purpose. The Receiving Party will engage with the Company in the role of "
        "{{ role_title }} starting on {{ start_date }}. In the course of that "
        "engagement the Receiving Party may receive information that the Company "
        "considers proprietary and confidential (“Confidential Information”)."
    ),
    (
        "2. Definition. Confidential Information includes, without limitation, "
        "business plans, source code, customer lists, financial records, product "
        "roadmaps, hiring data, internal documents, and any other non-public "
        "information disclosed in writing, orally, or through observation."
    ),
    (
        "3. Obligations. The Receiving Party will (a) hold the Confidential "
        "Information in strict confidence, (b) use it solely for the purpose of "
        "performing duties for the Company, (c) not disclose it to any third "
        "party without prior written consent, and (d) take reasonable measures "
        "to protect it from unauthorised access."
    ),
    (
        "4. Exclusions. This Agreement does not apply to information that "
        "(a) was publicly known at the time of disclosure, (b) becomes publicly "
        "known after disclosure through no fault of the Receiving Party, "
        "(c) was already in the Receiving Party’s possession without "
        "confidentiality restrictions before disclosure, or (d) must be "
        "disclosed pursuant to a binding court order or law."
    ),
    (
        "5. Term. The obligations under this Agreement remain in effect during "
        "the engagement and for a period of three (3) years following the end "
        "of the engagement, regardless of the reason for termination."
    ),
    (
        "6. Return of Materials. Upon termination of the engagement, or earlier "
        "upon written request, the Receiving Party will return or destroy all "
        "Confidential Information in their possession and certify the same in "
        "writing if requested."
    ),
    (
        "7. No License. Nothing in this Agreement grants the Receiving Party "
        "any right, title, or licence in or to the Confidential Information "
        "other than the limited right to use it for the agreed purpose."
    ),
    (
        "8. Remedies. The Receiving Party acknowledges that any breach of this "
        "Agreement may cause the Company irreparable harm for which monetary "
        "damages would be inadequate, and that the Company is entitled to seek "
        "injunctive relief in addition to any other remedies available at law."
    ),
    (
        "9. Governing Law. This Agreement is governed by and construed in "
        "accordance with the laws of {{ jurisdiction }}, without regard to its "
        "conflict-of-laws principles. Any dispute arising under this Agreement "
        "will be resolved by the courts located in {{ jurisdiction }}."
    ),
    (
        "10. Entire Agreement. This Agreement constitutes the entire "
        "understanding between the parties on this subject and supersedes all "
        "prior negotiations, representations, and agreements. Any modification "
        "must be in writing and signed by both parties."
    ),
    (
        "By signing below, the parties confirm they have read, understood, and "
        "agreed to be bound by the terms of this Agreement."
    ),
]


def _build_nda_docx() -> bytes:
    """Build the default NDA as a fresh DOCX in-memory.

    Synchronous — keep at module level. Called via asyncio.to_thread.
    """
    try:
        from docx import Document  # type: ignore[import-not-found]
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx not installed but required for default NDA template."
        ) from exc

    doc = Document()
    # Set a readable base font.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(_NDA_TITLE, level=0)
    for run in title.runs:
        run.font.size = Pt(20)

    for para in _NDA_PARAGRAPHS:
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(8)

    # Signature block.
    doc.add_paragraph()
    doc.add_paragraph(
        "For the Company:                                                 "
        "For the Receiving Party:"
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "_______________________                                "
        "_______________________"
    )
    doc.add_paragraph("Name: {{ company_name }}")
    doc.add_paragraph("Name: {{ candidate_name }}")
    doc.add_paragraph("Title: Authorised Signatory")
    doc.add_paragraph("Date: {{ today_date }}")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


async def get_default_template_docx(template_kind: str) -> bytes | None:
    """Return the bytes for a default template if we ship one for this
    kind, else None. The agent treats a None as "no fallback — block run
    on missing-template"."""
    if template_kind == "nda":
        return await asyncio.to_thread(_build_nda_docx)
    return None
