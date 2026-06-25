"""RFP Response service (#22 — Sales Enablement, second half).

Three-stage pipeline:

  1. extract_requirements()
     PDF/DOCX upload → reuse `ingestion.parser.parse_document` to extract text
     → Gemini structures the prose into a list of discrete requirements.
     Persists with status='extracted' for the rep to review/edit.

  2. (rep edits requirements in UI, calls PATCH endpoint to update the
     `requirements` array)

  3. generate_response()
     For each requirement, hybrid_search the KB. Above threshold → "matched"
     with the best source doc + a synthesized response paragraph.
     Below → "gap" with a flag for legal review.
     Assembles a .docx via python-docx with visually-distinct gap sections,
     uploads to Supabase Storage, returns a signed URL.

`python-docx` is already in pyproject (line 21).
"""
from __future__ import annotations

import asyncio
import logging
import os
import tempfile
import uuid
from datetime import UTC, datetime
from typing import Any

from app.database import get_service_client
from app.models.sales_enablement import RfpMatch, RfpRequirement
from app.services.agents.kb_synthesis import synthesize_json, synthesize_text
from app.services.ingestion.parser import parse_document
from app.services.retrieval.hybrid_search import hybrid_search

log = logging.getLogger(__name__)


_RFP_EXTRACT_SYSTEM = """You parse an RFP (Request for Proposal) document into discrete requirements.

A "requirement" is one specific thing the customer is asking us to commit to or describe — capabilities, certifications, SLAs, integration support, pricing breakdowns, security controls, references.

Rules:
- Split umbrella sections into atomic requirements. "Describe your security posture" is too broad; produce one item per testable sub-claim (SOC 2 status, data residency, encryption, access controls).
- Keep customer language. Do not paraphrase requirements into our preferred terminology.
- Categorize each: security, integration, pricing, support, capability, compliance, references, other.

Output JSON only:
{
  "requirements": [
    {"id": "R1", "requirement_text": "string", "category": "string"}
  ]
}
""".strip()


_RFP_MATCH_SYSTEM = """You draft a single RFP-response paragraph for a customer requirement, grounded in the provided context. Do NOT make up facts not present in the context. If the context is too thin, say so plainly so legal can review.

Output one paragraph, 60–180 words. No headers. No JSON.
""".strip()


# ── Stage 1: extract ────────────────────────────────────────────────────────


async def extract_requirements(
    *,
    org_id: str,
    user_id: str,
    file_bytes: bytes,
    filename: str,
) -> dict[str, Any]:
    """Upload + parse + structure. Returns the created row."""
    ext = (filename.rsplit(".", 1)[-1] or "").lower()
    if ext not in {"pdf", "docx", "doc", "txt", "md"}:
        raise ValueError(f"unsupported_rfp_format: .{ext}")

    file_type = {"pdf": "pdf", "docx": "docx", "doc": "docx", "txt": "txt", "md": "md"}[ext]
    try:
        segments = await asyncio.to_thread(parse_document, file_bytes, file_type)
    except Exception as exc:
        raise RuntimeError(f"rfp_parse_failed: {exc}") from exc

    full_text = "\n\n".join(seg.text for seg in segments if (seg.text or "").strip())
    if not full_text.strip():
        raise RuntimeError("rfp_empty_document")

    # Trim aggressively for Gemini — 60k chars is plenty for a typical RFP
    # and keeps us well inside Flash's input cap.
    extract_text = full_text[:60000]

    try:
        result = await synthesize_json(
            system_prompt=_RFP_EXTRACT_SYSTEM,
            user_prompt=f"## RFP source\n\n{extract_text}",
            temperature=0.0,  # extraction is deterministic — no creativity needed
            timeout=120.0,
        )
    except Exception as exc:
        raise RuntimeError(f"rfp_extraction_failed: {exc}") from exc

    raw = (result or {}).get("requirements") if isinstance(result, dict) else None
    if not raw or not isinstance(raw, list):
        raise RuntimeError("rfp_extraction_returned_no_requirements")

    # Validate via Pydantic; drop malformed entries rather than fail the whole
    # extraction — the rep can still edit/add manually.
    requirements: list[RfpRequirement] = []
    for i, r in enumerate(raw):
        try:
            requirements.append(
                RfpRequirement(
                    id=r.get("id") or f"R{i+1}",
                    requirement_text=r["requirement_text"],
                    category=r.get("category"),
                )
            )
        except Exception:
            continue

    if not requirements:
        raise RuntimeError("rfp_extraction_no_valid_requirements")

    svc = get_service_client()

    def _insert() -> dict[str, Any]:
        res = (
            svc.table("rfp_responses")
            .insert(
                {
                    "org_id": org_id,
                    "created_by": user_id,
                    "source_filename": filename,
                    "source_text": full_text[:200_000],  # debug/audit
                    "requirements": [r.model_dump() for r in requirements],
                    "status": "extracted",
                }
            )
            .execute()
        )
        rows = res.data or []
        if not rows:
            raise RuntimeError("rfp_insert_returned_no_rows")
        return rows[0]

    return await asyncio.to_thread(_insert)


# ── Stage 2: update requirements (no-op here, router calls a direct UPDATE) ──


# ── Stage 3: match + generate docx ──────────────────────────────────────────


_MATCH_THRESHOLD = 0.55  # cosine similarity below this → flag as gap


async def _match_one(
    *,
    org_id: str,
    req: RfpRequirement,
) -> RfpMatch:
    svc = get_service_client()
    try:
        hits = await hybrid_search(
            query=req.requirement_text, org_id=org_id, client=svc, k=4
        )
    except Exception as exc:
        log.warning("rfp.match.search_failed req=%s err=%s", req.id, exc)
        return RfpMatch(
            requirement_id=req.id,
            status="gap",
            flag_message="Search failed; review manually.",
        )

    top = hits[0] if hits else None
    similarity = (top.vector_similarity or top.similarity) if top else 0.0

    if not top or similarity < _MATCH_THRESHOLD:
        return RfpMatch(
            requirement_id=req.id,
            status="gap",
            flag_message=(
                f"No supporting documentation found above similarity {_MATCH_THRESHOLD:.2f}. "
                "Flagged for legal/SME review."
            ),
            confidence=similarity,
        )

    # Assemble a tight context window from top 3 hits.
    context = "\n\n---\n\n".join(
        f"[{h.document_name}] {(h.content or '').strip()[:1500]}" for h in hits[:3]
    )
    prompt = (
        f"## Customer requirement\n{req.requirement_text}\n\n"
        f"## Our internal documentation\n{context}\n\n"
        "## Output\nWrite one customer-facing response paragraph."
    )
    try:
        response_text = await synthesize_text(
            system_prompt=_RFP_MATCH_SYSTEM,
            user_prompt=prompt,
            temperature=0.2,
            timeout=45.0,
        )
    except Exception as exc:
        log.warning("rfp.match.synthesize_failed req=%s err=%s", req.id, exc)
        return RfpMatch(
            requirement_id=req.id,
            status="gap",
            flag_message="AI failed to draft a response; review manually.",
            confidence=similarity,
        )

    return RfpMatch(
        requirement_id=req.id,
        status="matched",
        source_doc=top.document_name,
        response_text=response_text,
        confidence=similarity,
    )


async def generate_response(
    *,
    org_id: str,
    rfp_id: str,
) -> dict[str, Any]:
    """Run matching for every requirement, build the .docx, persist, return row."""
    svc = get_service_client()

    def _fetch() -> dict[str, Any] | None:
        res = (
            svc.table("rfp_responses")
            .select("*")
            .eq("id", rfp_id)
            .eq("org_id", org_id)
            .maybe_single()
            .execute()
        )
        return res.data if res else None

    row = await asyncio.to_thread(_fetch)
    if not row:
        raise LookupError("rfp_not_found")

    requirements = [RfpRequirement(**r) for r in row.get("requirements") or []]
    if not requirements:
        raise ValueError("rfp_no_requirements")

    await asyncio.to_thread(
        lambda: svc.table("rfp_responses")
        .update({"status": "generating", "error_message": None})
        .eq("id", rfp_id)
        .execute()
    )

    try:
        # Cap concurrency at 5 to avoid blowing the LLM rate limit on a 50-
        # requirement RFP. asyncio.Semaphore wraps each task.
        sem = asyncio.Semaphore(5)

        async def _bounded(req: RfpRequirement) -> RfpMatch:
            async with sem:
                return await _match_one(org_id=org_id, req=req)

        matches: list[RfpMatch] = await asyncio.gather(
            *[_bounded(r) for r in requirements]
        )

        gap_count = sum(1 for m in matches if m.status == "gap")
        docx_path = await asyncio.to_thread(
            _build_docx,
            requirements=requirements,
            matches=matches,
            company_name=(row.get("source_filename") or "RFP").rsplit(".", 1)[0],
        )

        def _finalize() -> dict[str, Any]:
            res = (
                svc.table("rfp_responses")
                .update(
                    {
                        "matches": [m.model_dump() for m in matches],
                        "status": "ready",
                        "output_file_path": docx_path,
                        "gap_count": gap_count,
                        "generated_at": datetime.now(UTC).isoformat(),
                    }
                )
                .eq("id", rfp_id)
                .execute()
            )
            return (res.data or [{}])[0]

        return await asyncio.to_thread(_finalize)
    except Exception as exc:
        log.exception("rfp.generate.failed rfp_id=%s", rfp_id)
        await asyncio.to_thread(
            lambda: svc.table("rfp_responses")
            .update({"status": "failed", "error_message": str(exc)})
            .eq("id", rfp_id)
            .execute()
        )
        raise


def _build_docx(
    *,
    requirements: list[RfpRequirement],
    matches: list[RfpMatch],
    company_name: str,
) -> str:
    """Synchronous python-docx assembly. Writes to a tempfile, returns the
    path. The router uploads to Supabase Storage and signs a URL — keeping
    the storage side-effect out of here lets us unit-test this purely.

    Visual scheme:
      * Title page with the source filename + generation timestamp.
      * Per-requirement section. Matched: black body text + a small grey
        "Source: <doc>" caption. Gap: highlighted in red + a callout box
        flagging it for legal review.
    """
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt, RGBColor

    doc = Document()

    title = doc.add_heading("RFP Response", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle = doc.add_paragraph(f"Source: {company_name}")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ts = doc.add_paragraph(
        f"Generated: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    ts.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()  # spacer

    # Gap summary up top so legal can find work without scrolling.
    gap_matches = [m for m in matches if m.status == "gap"]
    if gap_matches:
        warn_heading = doc.add_heading(
            f"⚠ {len(gap_matches)} requirement(s) flagged for legal review",
            level=1,
        )
        for run in warn_heading.runs:
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        for m in gap_matches:
            req_text = next(
                (r.requirement_text for r in requirements if r.id == m.requirement_id),
                m.requirement_id,
            )
            p = doc.add_paragraph()
            p.add_run(f"  • [{m.requirement_id}] ").bold = True
            p.add_run(req_text[:200])
        doc.add_page_break()

    # Per-requirement sections.
    by_id = {m.requirement_id: m for m in matches}
    for r in requirements:
        m = by_id.get(r.id)
        h = doc.add_heading(f"{r.id}. {r.requirement_text}", level=2)
        if r.category:
            cat = doc.add_paragraph()
            cat_run = cat.add_run(f"Category: {r.category}")
            cat_run.italic = True
            cat_run.font.size = Pt(9)
            cat_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if m and m.status == "matched":
            body = doc.add_paragraph(m.response_text or "")
            if m.source_doc:
                src = doc.add_paragraph()
                src_run = src.add_run(f"Source: {m.source_doc}")
                src_run.italic = True
                src_run.font.size = Pt(9)
                src_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        else:
            gap_p = doc.add_paragraph()
            gap_run = gap_p.add_run(
                "⚠ GAP — flagged for legal review. "
                + (m.flag_message if m else "No documentation matched.")
            )
            gap_run.bold = True
            gap_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    out_path = os.path.join(
        tempfile.gettempdir(), f"rfp_response_{uuid.uuid4().hex}.docx"
    )
    doc.save(out_path)
    return out_path
