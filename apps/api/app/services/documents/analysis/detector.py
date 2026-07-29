"""Locate the fill-points in an HR-authored document.

Detection is deterministic and local — no LLM — and is deliberately biased
toward RECALL. Over-flagging is cheap: a false positive is one click for HR to
reject, and rejected rows are remembered so re-analysis never re-proposes them.
Under-flagging is what ships a contract with a blank signature line, so a missed
fill-point is the expensive failure.

Why this exists rather than the previous `blank_detector`
--------------------------------------------------------
The old detector could only see fill-points that were *empty*: a run of
underscores, a bracketed label, a label with nothing after it, an empty table
cell. That is a correct model of a blank form, and a completely wrong model of
what HR actually uploads.

The common case — and the one in the product spec — is a **finished document**
used as the sample:

    Dear Rahul,
    We are pleased to offer you the position of Software Engineer.
    Joining Date: 15 August 2026
    Annual Salary: ₹12,00,000

Not one of those values matches an empty-shape heuristic, so the old pipeline
detected nothing at all and HR concluded the analysis was broken. It was: the
detector was answering "where are the gaps?" when the question is "which text
changes per candidate?".

So detection now runs two families of heuristic:

  **Empty fill-points** (carried over, they were correct)
    1. `blank_marker`      — `________`, `[NAME]`, `<NAME>`, `....`, `{name}`
    2. `label_colon_empty` — `Employee Signature:` and then nothing
    3. `empty_cell`        — a blank cell beside a labelled one

  **Occupied fill-points** (new — the recall gap)
    4. `label_colon_value` — `Joining Date: 15 August 2026`, span over the value
    5. `table_cell_value`  — `Designation | Software Engineer`, span over the value
    6. `pattern_value`     — a standalone amount, date, email, or phone

A seventh kind, `llm_literal`, is produced by `analyzer.py` for values with no
label at all (the `Rahul` in `Dear Rahul,`). It is listed in `DETECTION_KINDS`
here because it shares the anchor contract, but it is never produced by this
module: the model proposes a substring, and the analyzer locates it by exact
search. Offsets in this system are *always* computed by `str.find`, never
authored by a model.

Every candidate is addressed by `(paragraph_index, start_offset, end_offset)` in
`canonical_paragraphs()` order, so detection, the persisted slot row, and the
renderer all point at the same characters.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Any

from docx import Document  # python-docx
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.services.documents.constants import (
    ACTION_INSERT_AFTER_LABEL,
    ACTION_INSERT_EMPTY_CELL,
    ACTION_REPLACE_SPAN,
)
from app.services.documents.docx_positions import (
    canonical_paragraphs,
    paragraph_run_text,
)

# ── Detection kinds ────────────────────────────────────────────────────────

KIND_BLANK_MARKER = "blank_marker"
KIND_LABEL_COLON_EMPTY = "label_colon_empty"
KIND_EMPTY_CELL = "empty_cell"
KIND_LABEL_COLON_VALUE = "label_colon_value"
KIND_TABLE_CELL_VALUE = "table_cell_value"
KIND_PATTERN_VALUE = "pattern_value"
KIND_LLM_LITERAL = "llm_literal"

DETECTION_KINDS = (
    KIND_BLANK_MARKER,
    KIND_LABEL_COLON_EMPTY,
    KIND_EMPTY_CELL,
    KIND_LABEL_COLON_VALUE,
    KIND_TABLE_CELL_VALUE,
    KIND_PATTERN_VALUE,
    KIND_LLM_LITERAL,
)

# Kinds where the document already contains a real value that a generation will
# OVERWRITE. The builder UI must show these differently: rejecting a
# `blank_marker` leaves an ugly blank, but rejecting one of these leaves the
# sample candidate's data in a document sent to someone else.
OCCUPIED_KINDS = frozenset({
    KIND_LABEL_COLON_VALUE,
    KIND_TABLE_CELL_VALUE,
    KIND_PATTERN_VALUE,
    KIND_LLM_LITERAL,
})


# ── Patterns ───────────────────────────────────────────────────────────────

# A blank fill-spot marked by visible characters, matched as a SPAN so we get
# exact offsets.
_BLANK_SPAN_RE = re.compile(
    r"_{3,}"                            # ________
    r"|\[[^\[\]{}\n]{1,60}\]"           # [CANDIDATE NAME]
    r"|<[^<>{}\n]{1,60}>"               # <NAME>
    r"|\.{3,}"                          # dotted leader ......
    r"|…"                               # unicode ellipsis
    r"|(?<!\{)\{[^{}\n]{1,60}\}(?!\})"  # {name}
    r"|\{\{[^{}\n]{1,60}\}\}"           # {{ name }} — a hand-authored placeholder
)

# A label with nothing after it. Capped at 60 chars so it fires on
# `Employee Signature:` but not on a prose sentence ending in a colon.
_LABEL_COLON_EMPTY_RE = re.compile(r"^[A-Z][A-Za-z0-9 /()&.'\-]{0,58}[:：]$")

# `Label: value` — the filled counterpart. The label group is validated further
# by `_is_field_label`; this only establishes the shape.
_LABEL_VALUE_RE = re.compile(r"^([A-Z][A-Za-z0-9 /()&.'\-]{0,58})[:：][ \t]*(\S.*)$")

# Colon-terminated prose lead-ins, not field labels.
_LEAD_IN_SUFFIXES = (
    "as follows", "the following", "following", "as under", "as below",
    "below", "for example", "namely", "provided that", "in witness whereof",
    "witnesseth", "whereas", "note", "important", "warning", "disclaimer",
)

# Words that only appear in a sentence, never in a form-field label. A single
# hit rejects. Deliberately excludes label filler like "of"/"the", because
# "Name of the Employee:" is a real label.
_PROSE_WORDS = frozenset({
    "are", "is", "was", "were", "be", "been", "being",
    "will", "shall", "may", "must", "should", "would", "can",
    "has", "have", "had", "do", "does", "did",
    "include", "includes", "including", "means", "agrees", "hereby",
    "you", "your", "we", "our", "they", "their", "this", "these", "those",
})

# Field labels common in HR documents. Used to (a) mark a table cell's
# neighbour as a fill-point even without a trailing colon and (b) raise
# confidence on a `Label: value` split. Covers the canonical candidate profile
# fields, not just the ones the previous fixed vocabulary happened to include.
_FIELD_LABEL_WORDS = frozenset({
    "name", "full name", "candidate name", "employee name", "employee",
    "father's name", "date of birth", "dob", "gender", "nationality",
    "date", "start date", "joining date", "date of joining", "effective date",
    "end date", "last date", "offer date", "issue date",
    "signature", "signed", "place", "witness", "for and on behalf of",
    "designation", "role", "title", "job title", "position", "grade", "level",
    "department", "team", "function", "business unit",
    "location", "work location", "office", "office location", "base location",
    "work mode", "employment type", "type of employment",
    "ctc", "salary", "compensation", "annual salary", "gross salary",
    "fixed pay", "variable pay", "bonus", "stipend", "remuneration",
    "basic", "hra", "allowance", "special allowance",
    "email", "email id", "e-mail", "phone", "contact", "mobile",
    "contact number", "phone number",
    "address", "current address", "permanent address", "correspondence address",
    "city", "state", "country", "pin", "pincode", "zip",
    "manager", "reporting manager", "reports to", "supervisor",
    "employee id", "id", "emp id", "employee code",
    "probation", "probation period", "notice period", "working hours",
    "jurisdiction", "governing law", "company", "company name", "employer",
    "hr", "hr contact", "hr manager",
})

# Standalone value patterns, used where there is no label to key off.
_CURRENCY_RE = re.compile(
    r"(?:₹|Rs\.?|INR|USD|US\$|\$|€|£|AED|SGD)\s?\d[\d, .]*"
    r"(?:\s?(?:lakhs?|lacs?|crores?|million|billion|mn|bn|k))?",
    re.IGNORECASE,
)

_MONTHS = (
    r"Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:t)?(?:ember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?"
)
_DATE_RE = re.compile(
    rf"\b\d{{1,2}}(?:st|nd|rd|th)?\s+(?:{_MONTHS})\.?,?\s+\d{{4}}\b"
    rf"|\b(?:{_MONTHS})\.?\s+\d{{1,2}}(?:st|nd|rd|th)?,?\s+\d{{4}}\b"
    r"|\b\d{4}-\d{2}-\d{2}\b"
    r"|\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
    re.IGNORECASE,
)

_EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b")

# Requires either a leading + or 10+ digits, so it does not swallow dates or
# short reference numbers.
_PHONE_RE = re.compile(r"\+\d[\d\s\-()]{7,}\d|\b\d{10,}\b")

# How much surrounding text to capture on each side — enough for the classifier
# (and HR) to read the label without bloating the prompt.
_CONTEXT_CHARS = 60

# Upper bound on candidates handed to the classifier. A 40-page policy manual
# uploaded by mistake should degrade to "too many fields, review manually"
# rather than a five-figure token bill.
MAX_CANDIDATES = 400


@dataclass(frozen=True)
class FillCandidate:
    """A fill-point located deterministically by exact position.

    Offsets are the single source of truth for both the classification prompt
    and the later substitution, so nothing can drift. Two identical-looking
    fill-points (two `________` signature lines) are distinguished by
    `paragraph_index` + `start_offset`, never by their identical text.

    For the two `insert_*` actions there is no span to replace:
    `start_offset == end_offset` marks an insertion point and `original_text`
    is empty.
    """

    candidate_id: int
    paragraph_index: int
    start_offset: int
    end_offset: int
    original_text: str
    context_before: str
    context_after: str
    paragraph_kind: str
    action: str
    detection_kind: str
    # The label governing this fill-point, when detection found one. Fed to the
    # classifier as a strong hint and shown to HR in the builder. Empty when the
    # candidate has no label (a bare amount in a clause, an LLM literal).
    label_hint: str = ""

    @property
    def is_occupied(self) -> bool:
        """True when a real value sits here today and will be overwritten."""
        return self.detection_kind in OCCUPIED_KINDS


# ── Label classification ───────────────────────────────────────────────────


def _looks_like_field_label(bare: str) -> bool:
    """True if `bare` (a colon-stripped, lowercased string) reads as a field
    label rather than prose.

    Three guards, all necessary. The shape test alone accepts
    `The terms are as follows:` — a clause that must never have a value
    appended. The suffix list catches common legal lead-ins; the prose-word
    check catches the rest by grammar, since a field label contains no verb.
    """
    bare = bare.strip().lower()
    if not bare:
        return False
    if any(bare.endswith(suffix) for suffix in _LEAD_IN_SUFFIXES):
        return False
    words = bare.replace("/", " ").split()
    if any(w.strip(".,()") in _PROSE_WORDS for w in words):
        return False
    return len(words) <= 6


def _is_known_field_label(bare: str) -> bool:
    """True if the label is one we recognise outright — a strong signal that
    whatever follows it is candidate data rather than boilerplate."""
    return bare.strip().rstrip(":：").strip().lower() in _FIELD_LABEL_WORDS


def _is_field_label_paragraph(text: str) -> bool:
    """True for a whole paragraph that is exactly a field label and a colon."""
    stripped = text.strip()
    if not _LABEL_COLON_EMPTY_RE.match(stripped):
        return False
    return _looks_like_field_label(stripped[:-1])


# ── Heuristic 1: visible blank markers ─────────────────────────────────────


def _blank_marker_candidates(
    canonical: list[tuple[Paragraph, str]],
    *,
    skip_indices: set[int],
) -> list[FillCandidate]:
    out: list[FillCandidate] = []
    for idx, (paragraph, kind) in enumerate(canonical):
        if idx in skip_indices:
            continue
        text = paragraph_run_text(paragraph)
        if not text:
            continue
        for m in _BLANK_SPAN_RE.finditer(text):
            start, end = m.start(), m.end()
            out.append(
                FillCandidate(
                    candidate_id=-1,
                    paragraph_index=idx,
                    start_offset=start,
                    end_offset=end,
                    original_text=m.group(0),
                    context_before=text[max(0, start - _CONTEXT_CHARS):start],
                    context_after=text[end:end + _CONTEXT_CHARS],
                    paragraph_kind=kind,
                    action=ACTION_REPLACE_SPAN,
                    detection_kind=KIND_BLANK_MARKER,
                    label_hint=_label_before(text, start),
                )
            )
    return out


def _label_before(text: str, start: int) -> str:
    """Extract a label sitting immediately before a span, if there is one.

    `Joining Date: ________` → "Joining Date". Looks only at the current
    paragraph and only back to the previous colon, so a long clause containing
    a blank yields no spurious label.
    """
    head = text[:start].rstrip()
    if not head.endswith((":", "：")):
        # Not a `label: blank` shape — but the whole head may still be a short
        # label-like fragment ("Name ________").
        candidate = head.strip()
        if candidate and _looks_like_field_label(candidate) and len(candidate) <= 60:
            return candidate
        return ""
    label = head[:-1].strip()
    # Only take the trailing fragment after any earlier sentence break.
    for sep in (". ", "; ", "•"):
        if sep in label:
            label = label.rsplit(sep, 1)[-1].strip()
    return label if _looks_like_field_label(label) else ""


# ── Heuristic 2: label with nothing after it ───────────────────────────────


def _label_colon_empty_candidates(
    canonical: list[tuple[Paragraph, str]],
    *,
    skip_indices: set[int],
) -> list[FillCandidate]:
    out: list[FillCandidate] = []
    for idx, (paragraph, kind) in enumerate(canonical):
        if idx in skip_indices:
            continue
        text = paragraph_run_text(paragraph)
        if not text or not _is_field_label_paragraph(text):
            continue
        if _BLANK_SPAN_RE.search(text):
            continue  # already a marker-span candidate
        offset = len(text)
        out.append(
            FillCandidate(
                candidate_id=-1,
                paragraph_index=idx,
                start_offset=offset,
                end_offset=offset,
                original_text="",
                context_before=text[max(0, offset - _CONTEXT_CHARS):offset],
                context_after="",
                paragraph_kind=kind,
                action=ACTION_INSERT_AFTER_LABEL,
                detection_kind=KIND_LABEL_COLON_EMPTY,
                label_hint=text.strip().rstrip(":：").strip(),
            )
        )
    return out


# ── Heuristic 4: label with a value after it (NEW) ────────────────────────


def _label_colon_value_candidates(
    canonical: list[tuple[Paragraph, str]],
    *,
    skip_indices: set[int],
) -> list[FillCandidate]:
    """`Joining Date: 15 August 2026` → a span over `15 August 2026`.

    This is the filled counterpart of heuristic 2 and the single biggest recall
    win over the previous detector, which required the value to be ABSENT.

    Guarded on both sides: the label must read as a field label, and the value
    must not itself be a blank marker (that case belongs to heuristic 1) nor a
    long clause (a colon inside prose is not a field).
    """
    out: list[FillCandidate] = []
    for idx, (paragraph, kind) in enumerate(canonical):
        if idx in skip_indices:
            continue
        text = paragraph_run_text(paragraph)
        if not text:
            continue
        match = _LABEL_VALUE_RE.match(text.strip())
        if not match:
            continue

        label, value = match.group(1).strip(), match.group(2).strip()
        if not _looks_like_field_label(label):
            continue
        if _BLANK_SPAN_RE.search(value):
            continue  # the "value" is a blank marker — heuristic 1 owns it
        # A value that runs on is a sentence, not a field. Known labels get more
        # rope because "Address: 12 MG Road, Bengaluru, Karnataka 560001" is
        # legitimately long.
        limit = 120 if _is_known_field_label(label) else 60
        if len(value) > limit or len(value.split()) > 14:
            continue

        # Locate the value in the ORIGINAL (unstripped) text so offsets are real.
        start = text.rfind(value)
        if start < 0:
            continue
        end = start + len(value)
        out.append(
            FillCandidate(
                candidate_id=-1,
                paragraph_index=idx,
                start_offset=start,
                end_offset=end,
                original_text=value,
                context_before=text[max(0, start - _CONTEXT_CHARS):start],
                context_after=text[end:end + _CONTEXT_CHARS],
                paragraph_kind=kind,
                action=ACTION_REPLACE_SPAN,
                detection_kind=KIND_LABEL_COLON_VALUE,
                label_hint=label,
            )
        )
    return out


# ── Table walking (heuristics 3 and 5) ────────────────────────────────────


def _iter_tables(container: Any) -> list[Table]:
    """All tables in a document, including tables nested inside cells."""
    found: list[Table] = []
    for table in getattr(container, "tables", []) or []:
        found.append(table)
        for row in table.rows:
            for cell in row.cells:
                found.extend(_iter_tables(cell))
    return found


def _cell_text(cell: _Cell) -> str:
    return " ".join(paragraph_run_text(p) for p in cell.paragraphs).strip()


def _is_cell_label(text: str) -> bool:
    bare = text.strip().rstrip(":：").strip().lower()
    if not bare or len(bare) > 60:
        return False
    return text.strip().endswith((":", "：")) or bare in _FIELD_LABEL_WORDS


def _table_candidates(
    doc: Any,
    canonical: list[tuple[Paragraph, str]],
) -> tuple[list[FillCandidate], set[int]]:
    """Walk `label | value` table rows.

    Produces `empty_cell` candidates where the value cell is blank and
    `table_cell_value` candidates where it already holds a value. Returns the
    candidates plus the canonical indices of every paragraph consumed (both
    label cells and value cells), so later heuristics don't produce a second,
    competing fill-point for the same field.
    """
    index_by_element = {
        paragraph._element: idx for idx, (paragraph, _kind) in enumerate(canonical)
    }

    out: list[FillCandidate] = []
    consumed: set[int] = set()

    for table in _iter_tables(doc):
        for row in table.rows:
            try:
                cells = list(row.cells)
            except (IndexError, ValueError):
                # Malformed / merged row geometry — skip rather than guess.
                continue
            for left, right in zip(cells, cells[1:]):
                if left._tc is right._tc:
                    continue  # horizontally merged: the same cell twice
                label_text = _cell_text(left)
                if not _is_cell_label(label_text):
                    continue
                value_text = _cell_text(right)

                target = next(iter(right.paragraphs), None)
                if target is None:
                    continue
                idx = index_by_element.get(target._element)
                if idx is None:
                    continue

                label = label_text.strip().rstrip(":：").strip()

                if not value_text:
                    out.append(
                        FillCandidate(
                            candidate_id=-1,
                            paragraph_index=idx,
                            start_offset=0,
                            end_offset=0,
                            original_text="",
                            context_before=label_text[-_CONTEXT_CHARS:],
                            context_after="",
                            paragraph_kind="table",
                            action=ACTION_INSERT_EMPTY_CELL,
                            detection_kind=KIND_EMPTY_CELL,
                            label_hint=label,
                        )
                    )
                else:
                    # Span the value paragraph's own run text. Using the
                    # paragraph (not the joined cell text) keeps offsets valid
                    # for multi-paragraph cells.
                    para_text = paragraph_run_text(target)
                    if not para_text.strip():
                        continue
                    if _BLANK_SPAN_RE.search(para_text):
                        continue  # a marker inside the cell — heuristic 1 owns it
                    stripped = para_text.strip()
                    start = para_text.find(stripped)
                    out.append(
                        FillCandidate(
                            candidate_id=-1,
                            paragraph_index=idx,
                            start_offset=start,
                            end_offset=start + len(stripped),
                            original_text=stripped,
                            context_before=label_text[-_CONTEXT_CHARS:],
                            context_after="",
                            paragraph_kind="table",
                            action=ACTION_REPLACE_SPAN,
                            detection_kind=KIND_TABLE_CELL_VALUE,
                            label_hint=label,
                        )
                    )

                consumed.add(idx)
                for p in left.paragraphs:
                    label_idx = index_by_element.get(p._element)
                    if label_idx is not None:
                        consumed.add(label_idx)

    return out, consumed


# ── Heuristic 6: standalone value patterns (NEW) ──────────────────────────


def _pattern_value_candidates(
    canonical: list[tuple[Paragraph, str]],
    *,
    skip_indices: set[int],
) -> list[FillCandidate]:
    """Amounts, dates, emails and phone numbers appearing anywhere.

    The safety net for values with no label at all — `...a gross salary of
    ₹12,00,000 per annum...` mid-clause. Restricted to patterns that are
    unambiguous on sight; free text like a candidate's name has no pattern and
    is left to the LLM literal pass.
    """
    out: list[FillCandidate] = []
    patterns = (
        (_EMAIL_RE, "email"),
        (_CURRENCY_RE, "amount"),
        (_DATE_RE, "date"),
        (_PHONE_RE, "phone"),
    )

    for idx, (paragraph, kind) in enumerate(canonical):
        if idx in skip_indices:
            continue
        text = paragraph_run_text(paragraph)
        if not text.strip():
            continue

        claimed: list[tuple[int, int]] = []
        for regex, hint in patterns:
            for m in regex.finditer(text):
                start, end = m.start(), m.end()
                if any(s < end and start < e for s, e in claimed):
                    continue  # an earlier, higher-priority pattern took this span
                claimed.append((start, end))
                value = m.group(0).strip()
                if not value:
                    continue
                # Re-derive offsets against the trimmed value so trailing
                # whitespace never lands inside the replaced span.
                real_start = start + m.group(0).find(value)
                out.append(
                    FillCandidate(
                        candidate_id=-1,
                        paragraph_index=idx,
                        start_offset=real_start,
                        end_offset=real_start + len(value),
                        original_text=value,
                        context_before=text[max(0, real_start - _CONTEXT_CHARS):real_start],
                        context_after=text[real_start + len(value):real_start + len(value) + _CONTEXT_CHARS],
                        paragraph_kind=kind,
                        action=ACTION_REPLACE_SPAN,
                        detection_kind=KIND_PATTERN_VALUE,
                        label_hint=_label_before(text, real_start) or hint,
                    )
                )
    return out


# ── Assembly ───────────────────────────────────────────────────────────────


def _drop_overlaps(candidates: list[FillCandidate]) -> list[FillCandidate]:
    """Keep the first candidate claiming any given character range.

    Input order encodes precedence, so callers concatenate in priority order.
    Insertion points (start == end) never conflict with each other and are only
    dropped if an identical anchor already exists — the unique index on
    (version_id, paragraph_index, start_offset, end_offset) would reject a
    duplicate at write time anyway, and failing here is cheaper.
    """
    kept: list[FillCandidate] = []
    spans_by_para: dict[int, list[tuple[int, int]]] = {}

    for c in candidates:
        spans = spans_by_para.setdefault(c.paragraph_index, [])
        if c.start_offset == c.end_offset:
            conflict = any(
                s == c.start_offset and e == c.end_offset for s, e in spans
            )
        else:
            conflict = any(
                s < c.end_offset and c.start_offset < e
                for s, e in spans
                if s != e  # an insertion point does not block a span
            )
        if conflict:
            continue
        spans.append((c.start_offset, c.end_offset))
        kept.append(c)

    return kept


def _renumber(candidates: list[FillCandidate]) -> list[FillCandidate]:
    """Assign sequential ids in document order.

    The classifier keys its response on `candidate_id`, so ids must be dense and
    stable within one analysis call.
    """
    ordered = sorted(candidates, key=lambda c: (c.paragraph_index, c.start_offset))
    return [
        FillCandidate(
            candidate_id=i,
            paragraph_index=c.paragraph_index,
            start_offset=c.start_offset,
            end_offset=c.end_offset,
            original_text=c.original_text,
            context_before=c.context_before,
            context_after=c.context_after,
            paragraph_kind=c.paragraph_kind,
            action=c.action,
            detection_kind=c.detection_kind,
            label_hint=c.label_hint,
        )
        for i, c in enumerate(ordered)
    ]


def find_all_candidates(
    docx_bytes: bytes,
    *,
    include_occupied: bool = True,
) -> list[FillCandidate]:
    """Run every heuristic over one parse of the document.

    Precedence, highest first — earlier heuristics claim their character ranges
    and later ones cannot re-claim them:

      1. table rows        (an explicit `label | value` grid is the strongest
                            structural signal in an HR document)
      2. blank markers     (a visible marker is unambiguous intent)
      3. label-colon-empty
      4. label-colon-value
      5. pattern values    (weakest — no structure, just shape)

    `include_occupied=False` restricts the result to empty fill-points, which is
    what a caller wants when re-analysing a document HR has already blanked out.
    """
    doc = Document(io.BytesIO(docx_bytes))
    canonical = canonical_paragraphs(doc)

    table_candidates, consumed = _table_candidates(doc, canonical)

    candidates: list[FillCandidate] = list(table_candidates)
    candidates += _blank_marker_candidates(canonical, skip_indices=consumed)
    candidates += _label_colon_empty_candidates(canonical, skip_indices=consumed)

    if include_occupied:
        candidates += _label_colon_value_candidates(canonical, skip_indices=consumed)
        candidates += _pattern_value_candidates(canonical, skip_indices=consumed)
    else:
        candidates = [c for c in candidates if not c.is_occupied]

    return _renumber(_drop_overlaps(candidates))


# ── Post-render safety net ─────────────────────────────────────────────────


def scan_for_unfilled_signals(text: str) -> list[str]:
    """Re-run the empty-shape signals over ALREADY-RENDERED text.

    Detection will never have perfect recall on arbitrary customer documents, so
    the last line of defence is checking the output: if a rendered offer letter
    still contains `________` or a bare `Signature:`, some fill-point was missed
    and HR should see that before the document reaches a candidate.

    Non-fatal by contract — the caller surfaces these as warnings, never as a
    render failure. A document with one missed field is still more useful to HR
    than no document at all.
    """
    warnings: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        match = _BLANK_SPAN_RE.search(line)
        if match:
            warnings.append(
                f"Still looks unfilled — '{match.group(0)}' in: {line[:80]}"
            )
            continue
        if _is_field_label_paragraph(line):
            warnings.append(f"Label with no value after it: {line[:80]}")
    return warnings
